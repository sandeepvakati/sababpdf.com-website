#!/usr/bin/env python3
"""
PDF to Word Converter - Table-Optimized Version
Preserves tables as editable Word tables with proper formatting (borders, colors)
Extracts text as editable content while keeping table structure intact
"""

import sys
import os
from pathlib import Path
import json
import logging
import shutil
from typing import List, Tuple, Optional, Dict
from io import BytesIO

# Configure logging
LOG_LEVEL_NAME = os.environ.get('PDF_TO_WORD_LOG_LEVEL', 'ERROR').upper()
LOG_LEVEL = getattr(logging, LOG_LEVEL_NAME, logging.ERROR)
logging.basicConfig(
    level=LOG_LEVEL,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Import required packages
FITZ_AVAILABLE = False
try:
    import fitz  # PyMuPDF
    FITZ_AVAILABLE = True
except ImportError as e:
    logger.error(f"PyMuPDF import error: {e}")

PDFPLUMBER_AVAILABLE = False
try:
    import pdfplumber
    import logging
    logging.getLogger('pdfminer').setLevel(logging.ERROR)
    logging.getLogger('pdfplumber').setLevel(logging.ERROR)
    PDFPLUMBER_AVAILABLE = True
except ImportError as e:
    logger.warning(f"pdfplumber import error: {e}")

DOCX_AVAILABLE = False
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm, Twips
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError as e:
    logger.error(f"python-docx import error: {e}")

PIL_AVAILABLE = False
try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError as e:
    logger.warning(f"Pillow import error: {e}")

# Constants
DEFAULT_DPI = 300
MARGIN_CM = 1.5


class PDFToWordTableConverter:
    """
    PDF to Word converter optimized for table preservation.
    Strategy: Extract tables as editable Word tables with borders/colors,
    extract text as editable paragraphs, preserve images.
    """

    def __init__(self, pdf_path: str, output_path: str, mode: str = 'editable'):
        self.pdf_path = pdf_path
        self.output_path = output_path
        self.mode = mode
        self.temp_dir = os.path.join(os.path.dirname(output_path), 'temp_conversion')
        self.pdf_doc = None
        self.page_count = 0
        self.pdfplumber_doc = None

    def open_pdf(self) -> bool:
        """Open PDF document with PyMuPDF"""
        if not FITZ_AVAILABLE:
            logger.error("PyMuPDF is required")
            return False

        try:
            self.pdf_doc = fitz.open(self.pdf_path)
            self.page_count = len(self.pdf_doc)
            logger.info(f"Opened PDF: {self.page_count} pages")
            return True
        except Exception as e:
            logger.error(f"Failed to open PDF: {e}")
            return False

    def open_pdfplumber(self) -> bool:
        """Open PDF with pdfplumber for table detection"""
        if not PDFPLUMBER_AVAILABLE:
            return False

        try:
            self.pdfplumber_doc = pdfplumber.open(self.pdf_path)
            logger.info("Opened pdfplumber document")
            return True
        except Exception as e:
            logger.warning(f"Failed to open pdfplumber: {e}")
            return False

    def close_pdf(self):
        """Close PDF documents"""
        if self.pdf_doc:
            self.pdf_doc.close()
            self.pdf_doc = None
        if self.pdfplumber_doc:
            self.pdfplumber_doc.close()
            self.pdfplumber_doc = None

    def cleanup_temp(self):
        """Clean up temporary files"""
        if os.path.exists(self.temp_dir):
            try:
                shutil.rmtree(self.temp_dir, ignore_errors=True)
            except Exception as e:
                logger.warning(f"Could not clean temp dir: {e}")

    def get_page_dimensions(self, page_num: int) -> Tuple[float, float]:
        """Get page dimensions in inches"""
        page = self.pdf_doc[page_num]
        rect = page.rect
        return rect.width / 72, rect.height / 72

    def extract_text_with_formatting(self, page_num: int, table_bboxes: List[Tuple] = None) -> List[Dict]:
        """
        Extract text blocks with formatting, excluding text that's inside tables.
        """
        if not FITZ_AVAILABLE:
            return []

        try:
            page = self.pdf_doc[page_num]
            text_dict = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)

            blocks = []
            for block in text_dict.get("blocks", []):
                if block.get("type") == 0:  # Text block
                    for line in block.get("lines", []):
                        spans = []
                        line_text = ""
                        for span in line.get("spans", []):
                            text = span.get("text", "")
                            line_text += text
                            if text.strip():
                                spans.append({
                                    'text': text,
                                    'font': span.get("font", "Calibri"),
                                    'size': span.get("size", 11),
                                    'color': span.get("color", 0),
                                    'bold': 'bold' in span.get("font", "").lower(),
                                    'italic': 'italic' in span.get("font", "").lower() or
                                             'oblique' in span.get("font", "").lower(),
                                })

                        if spans:
                            bbox = line.get("bbox")
                            # Check if this text block is inside any table region
                            is_in_table = False
                            if table_bboxes and bbox:
                                for table_bbox in table_bboxes:
                                    if self._bbox_overlaps(bbox, table_bbox):
                                        is_in_table = True
                                        break

                            if not is_in_table:
                                blocks.append({
                                    'type': 'text',
                                    'spans': spans,
                                    'bbox': bbox,
                                    'y_position': bbox[1] if bbox else 0,
                                })

            # Sort by y-position
            blocks.sort(key=lambda b: b['y_position'])
            return blocks
        except Exception as e:
            logger.error(f"Error extracting text from page {page_num}: {e}")
            return []

    def extract_tables(self, page_num: int) -> List[Dict]:
        """Extract tables using pdfplumber with enhanced detection"""
        if not PDFPLUMBER_AVAILABLE or not self.pdfplumber_doc:
            return []

        tables = []

        try:
            if page_num < len(self.pdfplumber_doc.pages):
                pdf_page = self.pdfplumber_doc.pages[page_num]
                
                # Try multiple table extraction strategies
                # Strategy 1: find_tables() with default settings
                pdf_tables = pdf_page.find_tables()
                
                for table_idx, table in enumerate(pdf_tables):
                    table_data = table.extract()
                    if table_data and len(table_data) > 0:
                        # Clean table data
                        max_cols = max(len(row) for row in table_data)
                        cleaned_data = []
                        for row in table_data:
                            cleaned_row = [str(cell).strip() if cell is not None else '' for cell in row]
                            while len(cleaned_row) < max_cols:
                                cleaned_row.append('')
                            cleaned_data.append(cleaned_row[:max_cols])

                        # Only add if table has meaningful content
                        if len(cleaned_data) >= 2:  # At least header + 1 row
                            tables.append({
                                'type': 'table',
                                'data': cleaned_data,
                                'bbox': table.bbox,
                                'index': table_idx,
                            })

                logger.info(f"Extracted {len(tables)} tables from page {page_num}")

        except Exception as e:
            logger.warning(f"Error extracting tables from page {page_num}: {e}")

        return tables

    def extract_images(self, page_num: int) -> List[Dict]:
        """Extract images from PDF page"""
        if not FITZ_AVAILABLE or not PIL_AVAILABLE:
            return []

        images = []
        os.makedirs(self.temp_dir, exist_ok=True)

        try:
            page = self.pdf_doc[page_num]
            image_list = page.get_images(full=True)

            for img_index, img_info in enumerate(image_list):
                try:
                    xref = img_info[0]
                    base_image = fitz.Pixmap(self.pdf_doc, xref)

                    if base_image.n - base_image.alpha > 3:  # CMYK
                        base_image = fitz.Pixmap(fitz.csRGB, base_image)

                    img_path = os.path.join(self.temp_dir, f'p{page_num}_img{img_index}.png')
                    base_image.save(img_path)

                    img_rect = img_info[1:5] if len(img_info) > 4 else None

                    images.append({
                        'type': 'image',
                        'path': img_path,
                        'bbox': img_rect,
                        'index': img_index,
                    })

                    base_image = None
                except Exception as e:
                    logger.warning(f"Could not extract image {img_index}: {e}")

        except Exception as e:
            logger.error(f"Error extracting images: {e}")

        return images

    def _bbox_overlaps(self, bbox1: Tuple, bbox2: Tuple, threshold: float = 0.5) -> bool:
        """Check if two bounding boxes overlap significantly"""
        if not bbox1 or not bbox2:
            return False

        x0_1, y0_1, x1_1, y1_1 = bbox1
        x0_2, y0_2, x1_2, y1_2 = bbox2

        # Calculate intersection area
        x0_i = max(x0_1, x0_2)
        y0_i = max(y0_1, y0_2)
        x1_i = min(x1_1, x1_2)
        y1_i = min(y1_1, y1_2)

        if x0_i >= x1_i or y0_i >= y1_i:
            return False

        intersection = (x1_i - x0_i) * (y1_i - y0_i)
        area1 = (x1_1 - x0_1) * (y1_1 - y0_1)
        
        overlap_ratio = intersection / area1 if area1 > 0 else 0
        return overlap_ratio > threshold

    def _set_table_style(self, table, header_bg_color: Tuple = (230, 245, 208)):
        """Set table borders and styling to match PDF appearance"""
        try:
            # Set table borders
            tbl = table._tbl
            tblPr = tbl.get_or_add_tblPr()
            tblBorders = OxmlElement('w:tblBorders')

            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '8')  # Border width
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), 'FF8C00')  # Orange border like in PDF
                tblBorders.append(border)

            tblPr.append(tblBorders)
        except Exception as e:
            logger.warning(f"Could not set table borders: {e}")

    def _add_table_to_doc(self, doc, table_info: Dict) -> None:
        """Add table to Word document with proper formatting"""
        try:
            table_data = table_info.get('data', [])
            if not table_data or len(table_data) < 2:
                return

            rows = len(table_data)
            cols = max(len(row) for row in table_data)

            # Add spacing
            doc.add_paragraph()

            # Create table
            table = doc.add_table(rows=rows, cols=cols)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Set table styling (borders)
            self._set_table_style(table)

            # Populate table cells
            for row_idx, row_data in enumerate(table_data):
                for col_idx, cell_text in enumerate(row_data):
                    if col_idx < cols:
                        cell = table.rows[row_idx].cells[col_idx]
                        cell.text = str(cell_text).strip() if cell_text else ''

                        # Format cell
                        for paragraph in cell.paragraphs:
                            paragraph.paragraph_format.space_before = Pt(2)
                            paragraph.paragraph_format.space_after = Pt(2)
                            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

                            for run in paragraph.runs:
                                run.font.size = Pt(10)
                                run.font.name = 'Calibri'

                                # Header row styling
                                if row_idx == 0:
                                    run.font.bold = True
                                    run.font.size = Pt(11)

                        # Header row background color
                        if row_idx == 0:
                            shading_elm = OxmlElement('w:shd')
                            shading_elm.set(qn('w:fill'), 'E6F5D0')  # Light green
                            cell._tc.get_or_add_tcPr().append(shading_elm)

                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            doc.add_paragraph()
            logger.info(f"Added table: {rows} rows x {cols} cols")

        except Exception as e:
            logger.error(f"Error adding table: {e}")

    def _add_text_block_to_doc(self, doc, block: Dict) -> None:
        """Add text block with formatting"""
        try:
            if not block.get('spans'):
                return

            paragraph = doc.add_paragraph()

            for span in block['spans']:
                run = paragraph.add_run(span['text'])

                font = run.font
                font.name = span['font'] or 'Calibri'
                font.size = Pt(span['size']) if span['size'] else Pt(11)

                if span.get('bold'):
                    font.bold = True
                if span.get('italic'):
                    font.italic = True

                # Apply color
                if span.get('color') and span['color'] != 0:
                    r = (span['color'] >> 16) & 0xFF
                    g = (span['color'] >> 8) & 0xFF
                    b = span['color'] & 0xFF
                    font.color.rgb = RGBColor(r, g, b)

        except Exception as e:
            logger.error(f"Error adding text block: {e}")

    def convert(self) -> Tuple[bool, str]:
        """Main conversion method"""
        if not os.path.exists(self.pdf_path):
            return False, f"PDF file not found: {self.pdf_path}"

        if not check_dependencies():
            return False, "Missing required dependencies"

        if not self.open_pdf():
            return False, "Failed to open PDF"

        try:
            os.makedirs(self.temp_dir, exist_ok=True)
            self.open_pdfplumber()

            # Create Word document
            doc = Document()

            # Set default style
            style = doc.styles['Normal']
            style.font.name = 'Calibri'
            style.font.size = Pt(11)
            style.paragraph_format.space_before = Pt(0)
            style.paragraph_format.space_after = Pt(6)
            style.paragraph_format.line_spacing = 1.15

            logger.info("Converting PDF to Word (table-optimized)...")

            for page_num in range(self.page_count):
                logger.info(f"Processing page {page_num + 1}/{self.page_count}")

                # Get page dimensions
                page_width_in, page_height_in = self.get_page_dimensions(page_num)

                # Add new page section
                if page_num > 0:
                    section = doc.add_section(WD_SECTION.NEW_PAGE)
                else:
                    section = doc.sections[0]

                # Set page size and margins
                section.page_width = Inches(page_width_in)
                section.page_height = Inches(page_height_in)
                section.top_margin = Cm(MARGIN_CM)
                section.bottom_margin = Cm(MARGIN_CM)
                section.left_margin = Cm(MARGIN_CM)
                section.right_margin = Cm(MARGIN_CM)

                # Extract tables first
                tables = self.extract_tables(page_num)
                table_bboxes = [t['bbox'] for t in tables if t.get('bbox')]

                # Extract text (excluding text inside tables)
                text_blocks = self.extract_text_with_formatting(page_num, table_bboxes)

                # Extract images
                images = self.extract_images(page_num)

                # Add content in order (simplified: text, then tables, then images)
                for block in text_blocks:
                    self._add_text_block_to_doc(doc, block)

                # Add tables
                for table_info in tables:
                    self._add_table_to_doc(doc, table_info)

                # Add images
                for img_info in images:
                    if os.path.exists(img_info['path']):
                        try:
                            paragraph = doc.add_paragraph()
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = paragraph.add_run()

                            with Image.open(img_info['path']) as img:
                                img_width_in = img.width / 150
                                img_height_in = img.height / 150

                                max_width = section.page_width.inches - 2 * MARGIN_CM
                                if img_width_in > max_width:
                                    scale = max_width / img_width_in
                                    img_width_in *= scale
                                    img_height_in *= scale

                                run.add_picture(img_info['path'],
                                              width=Inches(img_width_in),
                                              height=Inches(img_height_in))
                        except Exception as e:
                            logger.warning(f"Could not add image: {e}")

            # Save document
            doc.save(self.output_path)

            if os.path.exists(self.output_path):
                file_size = os.path.getsize(self.output_path)
                logger.info(f"Saved Word document: {file_size} bytes")
                return True, "Successfully converted to Word"

            return False, "Failed to create Word document"

        except Exception as e:
            logger.error(f"Conversion error: {e}")
            import traceback
            logger.error(traceback.format_exc())
            return False, f"Conversion error: {e}"

        finally:
            self.close_pdf()
            self.cleanup_temp()


def check_dependencies():
    """Check required dependencies"""
    required = {
        'PyMuPDF': FITZ_AVAILABLE,
        'python-docx': DOCX_AVAILABLE,
        'pdfplumber': PDFPLUMBER_AVAILABLE,
    }

    missing = [name for name, available in required.items() if not available]
    if missing:
        logger.error(f"Missing: {', '.join(missing)}")
        logger.error("Install: pip install pymupdf python-docx pdfplumber Pillow")
        return False
    return True


def main():
    """CLI entry point"""
    if len(sys.argv) < 3:
        print("Usage: python pdf_to_word_table.py <input_pdf> <output_docx> [mode]")
        sys.exit(1)

    input_pdf = sys.argv[1]
    output_docx = sys.argv[2]
    mode = sys.argv[3] if len(sys.argv) > 3 else 'editable'

    converter = PDFToWordTableConverter(input_pdf, output_docx, mode)
    success, message = converter.convert()

    result = {
        'success': success,
        'message': message,
        'input': input_pdf,
        'output': output_docx,
        'mode': mode,
    }

    print(json.dumps(result))
    sys.exit(0 if success else 1)


if __name__ == '__main__':
    main()
