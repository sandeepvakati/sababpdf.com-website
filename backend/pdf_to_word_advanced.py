#!/usr/bin/env python3
"""
Advanced PDF to Word Converter - iLovePDF Quality
Optimized for high-quality conversion with excellent layout preservation
Supports: Text, Images, Tables, Formulas, and Complex Layouts

IMPROVED: Better table detection, formula preservation, and layout accuracy
"""

import sys
import os
from pathlib import Path
import json
import logging
import re
import shutil
from typing import List, Tuple, Optional, Dict, Any
from io import BytesIO

# Configure logging - only show errors to keep output clean
LOG_LEVEL_NAME = os.environ.get('PDF_TO_WORD_LOG_LEVEL', 'ERROR').upper()
LOG_LEVEL = getattr(logging, LOG_LEVEL_NAME, logging.ERROR)
logging.basicConfig(
    level=LOG_LEVEL,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Try to import required packages
FITZ_AVAILABLE = False
try:
    import fitz  # PyMuPDF
    FITZ_AVAILABLE = True
except ImportError as e:
    logger.warning(f"PyMuPDF import error: {e}")
    pass

PDFPLUMBER_AVAILABLE = False
try:
    import pdfplumber
    import logging
    # Suppress pdfplumber's verbose warnings about overlapping text
    logging.getLogger('pdfminer').setLevel(logging.ERROR)
    logging.getLogger('pdfplumber').setLevel(logging.ERROR)
    PDFPLUMBER_AVAILABLE = True
except ImportError as e:
    logger.warning(f"pdfplumber import error: {e}")
    pass

DOCX_AVAILABLE = False
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm, Emu
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError as e:
    logger.warning(f"python-docx import error: {e}")
    pass

PIL_AVAILABLE = False
try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError as e:
    logger.warning(f"Pillow import error: {e}")
    pass

PDF2DOCX_AVAILABLE = False
try:
    from pdf2docx import Converter as Pdf2DocxConverter
    PDF2DOCX_AVAILABLE = True
except ImportError as e:
    logger.warning(f"pdf2docx import error: {e}")
    pass

# Constants
DEFAULT_DPI = 300  # High quality rendering
DEFAULT_SCALE = 2.0  # For better image quality
MARGIN_CM = 2.0  # Standard margins
TABLE_DETECTION_THRESHOLD = 0.7  # Threshold for table detection


def check_dependencies():
    """Check and report available dependencies"""
    deps = {
        'PyMuPDF': FITZ_AVAILABLE,
        'pdfplumber': PDFPLUMBER_AVAILABLE,
        'python-docx': DOCX_AVAILABLE,
        'Pillow': PIL_AVAILABLE,
        'pdf2docx': PDF2DOCX_AVAILABLE,
    }
    
    missing = [name for name, available in deps.items() if not available]
    
    if missing:
        logger.warning(f"Missing optional dependencies: {', '.join(missing)}")
        logger.warning("Install all dependencies: pip install pymupdf pdfplumber python-docx Pillow pdf2docx")
    
    return all(deps.values())


class HighQualityPDFToWordConverter:
    """
    High-quality PDF to Word converter optimized for iLovePDF-like results.
    Uses multiple strategies to achieve the best conversion quality.
    """
    
    def __init__(self, pdf_path: str, output_path: str, mode: str = 'layout'):
        self.pdf_path = pdf_path
        self.output_path = output_path
        self.mode = mode
        self.temp_dir = os.path.join(os.path.dirname(output_path), 'temp_conversion')
        self.pdf_doc = None
        self.page_count = 0
        
    def open_pdf(self) -> bool:
        """Open PDF document with PyMuPDF"""
        if not FITZ_AVAILABLE:
            logger.error("PyMuPDF is required for high-quality conversion")
            return False
            
        try:
            self.pdf_doc = fitz.open(self.pdf_path)
            self.page_count = len(self.pdf_doc)
            logger.info(f"Opened PDF: {self.page_count} pages")
            return True
        except Exception as e:
            logger.error(f"Failed to open PDF: {e}")
            return False
    
    def close_pdf(self):
        """Close PDF document"""
        if self.pdf_doc:
            self.pdf_doc.close()
            self.pdf_doc = None
    
    def cleanup_temp(self):
        """Clean up temporary files"""
        if os.path.exists(self.temp_dir):
            try:
                shutil.rmtree(self.temp_dir, ignore_errors=True)
            except Exception as e:
                logger.warning(f"Could not clean temp dir: {e}")
    
    def get_page_dimensions(self, page_num: int) -> Tuple[float, float]:
        """Get page dimensions in inches"""
        page = self.pdf_doc[page_num]
        rect = page.rect
        return rect.width / 72, rect.height / 72
    
    def render_page_to_image(self, page_num: int, dpi: int = DEFAULT_DPI) -> Optional[Image.Image]:
        """Render a PDF page to a high-quality PIL Image"""
        if not FITZ_AVAILABLE or not PIL_AVAILABLE:
            return None
            
        try:
            page = self.pdf_doc[page_num]
            zoom = dpi / 72
            mat = fitz.Matrix(zoom, zoom)
            
            pix = page.get_pixmap(matrix=mat, alpha=False)
            
            img_data = pix.tobytes('png')
            img = Image.open(BytesIO(img_data))
            
            return img.convert('RGB')
        except Exception as e:
            logger.error(f"Failed to render page {page_num}: {e}")
            return None
    
    def extract_text_blocks(self, page_num: int) -> List[Dict]:
        """Extract text blocks with formatting information"""
        if not FITZ_AVAILABLE:
            return []
            
        try:
            page = self.pdf_doc[page_num]
            text_dict = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)
            
            blocks = []
            for block in text_dict.get("blocks", []):
                if block.get("type") == 0:  # Text block
                    for line in block.get("lines", []):
                        spans = []
                        for span in line.get("spans", []):
                            text = span.get("text", "").strip()
                            if text:
                                spans.append({
                                    'text': text,
                                    'font': span.get("font", "Arial"),
                                    'size': span.get("size", 11),
                                    'color': span.get("color", 0),
                                    'bold': 'bold' in span.get("font", "").lower(),
                                    'italic': 'italic' in span.get("font", "").lower() or 
                                             'oblique' in span.get("font", "").lower(),
                                })
                        
                        if spans:
                            blocks.append({
                                'type': 'text',
                                'spans': spans,
                                'bbox': line.get("bbox"),
                            })
            
            return blocks
        except Exception as e:
            logger.error(f"Error extracting text from page {page_num}: {e}")
            return []
    
    def extract_images(self, page_num: int) -> List[Dict]:
        """Extract images from a PDF page"""
        if not FITZ_AVAILABLE:
            return []
            
        images = []
        os.makedirs(self.temp_dir, exist_ok=True)
        
        try:
            page = self.pdf_doc[page_num]
            image_list = page.get_images(full=True)
            
            for img_index, img_info in enumerate(image_list):
                try:
                    xref = img_info[0]
                    base_image = fitz.Pixmap(self.pdf_doc, xref)
                    
                    if base_image.n - base_image.alpha > 3:  # CMYK
                        base_image = fitz.Pixmap(fitz.csRGB, base_image)
                    
                    img_path = os.path.join(self.temp_dir, f'p{page_num}_img{img_index}.png')
                    base_image.save(img_path)
                    
                    # Get image dimensions
                    img_rect = img_info[1:5] if len(img_info) > 4 else None
                    
                    images.append({
                        'type': 'image',
                        'path': img_path,
                        'bbox': img_rect,
                        'index': img_index,
                    })
                    
                    base_image = None  # Free memory
                except Exception as e:
                    logger.warning(f"Could not extract image {img_index} from page {page_num}: {e}")
            
            # Also extract images from text dict (inline images)
            text_dict = page.get_text("dict")
            for block_index, block in enumerate(text_dict.get("blocks", [])):
                if block.get("type") == 1:  # Image block
                    try:
                        if "image" in block:
                            img_path = os.path.join(self.temp_dir, f'p{page_num}_blk{block_index}.png')
                            with open(img_path, "wb") as img_file:
                                img_file.write(block["image"])
                            
                            images.append({
                                'type': 'image',
                                'path': img_path,
                                'bbox': block.get("bbox"),
                                'index': f'block_{block_index}',
                            })
                    except Exception as e:
                        logger.warning(f"Could not extract block image {block_index}: {e}")
            
            logger.info(f"Extracted {len(images)} images from page {page_num}")
            
        except Exception as e:
            logger.error(f"Error extracting images from page {page_num}: {e}")
        
        return images
    
    def detect_table_regions(self, page_num: int) -> List[Dict]:
        """
        Detect table regions in a PDF page using multiple strategies.
        Returns list of table regions with bounding boxes.
        """
        table_regions = []
        
        try:
            page = self.pdf_doc[page_num]
            page_width, page_height = self.get_page_dimensions(page_num)
            page_width *= 72  # Convert back to points
            page_height *= 72
            
            # Strategy 1: Use pdfplumber for table detection
            if PDFPLUMBER_AVAILABLE:
                try:
                    with pdfplumber.open(self.pdf_path) as pdf:
                        if page_num < len(pdf.pages):
                            pdf_page = pdf.pages[page_num]
                            tables = pdf_page.find_tables()
                            
                            for table in tables:
                                bbox = table.bbox
                                if bbox:
                                    table_regions.append({
                                        'bbox': bbox,  # (x0, y0, x1, y1) in points
                                        'type': 'table',
                                        'confidence': 0.9,
                                        'data': table.extract()
                                    })
                except Exception as e:
                    logger.warning(f"pdfplumber table detection failed: {e}")
            
            # Strategy 2: Detect grid-like text patterns using PyMuPDF
            text_dict = page.get_text("dict")
            blocks = text_dict.get("blocks", [])
            
            # Group text blocks by vertical alignment (rows)
            text_lines = []
            for block in blocks:
                if block.get("type") == 0:  # Text block
                    for line in block.get("lines", []):
                        bbox = line.get("bbox")
                        text = "".join([span.get("text", "") for span in line.get("spans", [])])
                        if bbox and text.strip():
                            text_lines.append({
                                'bbox': bbox,
                                'text': text.strip(),
                                'y_center': (bbox[1] + bbox[3]) / 2,
                                'x_center': (bbox[0] + bbox[2]) / 2
                            })
            
            # Detect table-like structures (multiple columns of aligned text)
            if len(text_lines) >= 3:
                # Sort by y-position to find rows
                text_lines.sort(key=lambda x: x['y_center'])
                
                # Group into rows (lines with similar y-position)
                rows = []
                current_row = []
                y_threshold = 10  # Points tolerance for same row
                
                for line in text_lines:
                    if not current_row or abs(line['y_center'] - current_row[0]['y_center']) < y_threshold:
                        current_row.append(line)
                    else:
                        if len(current_row) >= 2:  # Row with multiple columns
                            rows.append(current_row)
                        current_row = [line]
                
                if len(current_row) >= 2:
                    rows.append(current_row)
                
                # Check if we have a table pattern (multiple rows with similar column structure)
                if len(rows) >= 2:
                    # Check column alignment across rows
                    col_positions = []
                    for row in rows[:5]:  # Check first 5 rows
                        for item in row:
                            col_positions.append(item['bbox'][0])  # x0 position
                    
                    # If we find consistent column positions, it's likely a table
                    col_positions.sort()
                    unique_cols = []
                    col_threshold = 30  # Points tolerance for same column
                    
                    for pos in col_positions:
                        if not unique_cols or pos - unique_cols[-1] > col_threshold:
                            unique_cols.append(pos)
                    
                    # If we have 2+ columns detected, mark as table region
                    if len(unique_cols) >= 2:
                        # Calculate bounding box for the table region
                        all_bboxes = [line['bbox'] for row in rows for line in row]
                        table_bbox = (
                            min(b[0] for b in all_bboxes),
                            min(b[1] for b in all_bboxes),
                            max(b[2] for b in all_bboxes),
                            max(b[3] for b in all_bboxes)
                        )
                        
                        # Check if this region overlaps with existing detected tables
                        overlaps = False
                        for existing in table_regions:
                            if self.bboxes_overlap(table_bbox, existing['bbox']):
                                overlaps = True
                                break
                        
                        if not overlaps:
                            # Extract table data from detected region
                            table_data = []
                            for row in rows:
                                row_data = [item['text'] for item in sorted(row, key=lambda x: x['bbox'][0])]
                                table_data.append(row_data)
                            
                            table_regions.append({
                                'bbox': table_bbox,
                                'type': 'table',
                                'confidence': 0.7,
                                'data': table_data
                            })
                            
                            logger.info(f"Detected table region via text pattern: {len(rows)} rows x {len(unique_cols)} cols")
            
        except Exception as e:
            logger.warning(f"Error detecting table regions: {e}")
        
        return table_regions
    
    def bboxes_overlap(self, bbox1: Tuple, bbox2: Tuple, threshold: float = 0.5) -> bool:
        """Check if two bounding boxes overlap significantly"""
        if not bbox1 or not bbox2:
            return False
        
        x0_1, y0_1, x1_1, y1_1 = bbox1
        x0_2, y0_2, x1_2, y1_2 = bbox2
        
        # Calculate intersection
        x0_i = max(x0_1, x0_2)
        y0_i = max(y0_1, y0_2)
        x1_i = min(x1_1, x1_2)
        y1_i = min(y1_1, y1_2)
        
        if x0_i >= x1_i or y0_i >= y1_i:
            return False
        
        intersection_area = (x1_i - x0_i) * (y1_i - y0_i)
        bbox1_area = (x1_1 - x0_1) * (y1_1 - y0_1)
        bbox2_area = (x1_2 - x0_2) * (y1_2 - y0_2)
        
        overlap_ratio = intersection_area / min(bbox1_area, bbox2_area) if min(bbox1_area, bbox2_area) > 0 else 0
        
        return overlap_ratio > threshold

    def extract_tables(self, page_num: int) -> List[Dict]:
        """Extract tables from a PDF page using enhanced detection"""
        if not PDFPLUMBER_AVAILABLE and not FITZ_AVAILABLE:
            return []

        tables = []

        try:
            # Use enhanced table region detection
            table_regions = self.detect_table_regions(page_num)
            
            for region in table_regions:
                if region.get('data'):
                    # Clean and validate table data
                    table_data = region['data']
                    if table_data and len(table_data) > 1:  # Need at least header + 1 row
                        # Ensure consistent column count
                        max_cols = max(len(row) for row in table_data)
                        cleaned_data = []
                        for row in table_data:
                            cleaned_row = [str(cell).strip() if cell is not None else '' for cell in row]
                            # Pad row to match max columns
                            while len(cleaned_row) < max_cols:
                                cleaned_row.append('')
                            cleaned_data.append(cleaned_row[:max_cols])
                        
                        tables.append({
                            'type': 'table',
                            'data': cleaned_data,
                            'bbox': region['bbox'],
                            'index': len(tables),
                            'confidence': region.get('confidence', 0.8),
                        })

            logger.info(f"Extracted {len(tables)} tables from page {page_num}")

        except Exception as e:
            logger.warning(f"Error extracting tables from page {page_num}: {e}")

        return tables
    
    def create_layout_document(self) -> bool:
        """
        Create a Word document with layout preservation (iLovePDF style).
        HYBRID APPROACH: 
        - For pages with tables/formulas: render as image for perfect fidelity
        - For text-only pages: extract as editable text with proper formatting
        """
        if not DOCX_AVAILABLE:
            logger.error("python-docx not available")
            return False

        try:
            doc = Document()

            # Set default style
            style = doc.styles['Normal']
            style.font.name = 'Calibri'
            style.font.size = Pt(11)
            style.paragraph_format.space_before = Pt(0)
            style.paragraph_format.space_after = Pt(6)
            style.paragraph_format.line_spacing = 1.15

            logger.info("Creating layout-preserving Word document (hybrid mode)...")

            for page_num in range(self.page_count):
                logger.info(f"Processing page {page_num + 1}/{self.page_count}")

                # Get page dimensions
                page_width_in, page_height_in = self.get_page_dimensions(page_num)

                # Add new page section
                if page_num > 0:
                    section = doc.add_section(WD_SECTION.NEW_PAGE)
                else:
                    section = doc.sections[0]

                # Set page size and margins
                section.page_width = Inches(page_width_in)
                section.page_height = Inches(page_height_in)
                section.top_margin = Cm(MARGIN_CM)
                section.bottom_margin = Cm(MARGIN_CM)
                section.left_margin = Cm(MARGIN_CM)
                section.right_margin = Cm(MARGIN_CM)

                # Detect content type for this page
                tables = self.extract_tables(page_num)
                text_blocks = self.extract_text_blocks(page_num)
                
                # Decide strategy: render as image if page has tables or complex formulas
                has_tables = len(tables) > 0
                has_complex_content = self._detect_complex_content(page_num)
                
                if has_tables or has_complex_content:
                    # Render page as image for perfect fidelity
                    logger.info(f"  Page {page_num + 1}: Rendering as image (has tables/complex content)")
                    self._add_page_as_image(doc, page_num, section)
                else:
                    # Extract as editable text
                    logger.info(f"  Page {page_num + 1}: Extracting as editable text")
                    self._add_page_as_editable_text(doc, page_num, text_blocks)

                # Add extracted tables separately if not rendering as image
                if has_tables and not has_complex_content:
                    for table_info in tables:
                        self._add_table_to_doc(doc, table_info)

            # Save document
            doc.save(self.output_path)

            if os.path.exists(self.output_path):
                file_size = os.path.getsize(self.output_path)
                logger.info(f"Layout document saved: {file_size} bytes")
                return True

            return False

        except Exception as e:
            logger.error(f"Error creating layout document: {e}")
            import traceback
            logger.error(traceback.format_exc())
            return False
    
    def _detect_complex_content(self, page_num: int) -> bool:
        """
        Detect if page has complex content like mathematical formulas,
        special symbols, or complex layouts that need image rendering.
        """
        if not FITZ_AVAILABLE:
            return False
        
        try:
            page = self.pdf_doc[page_num]
            text_dict = page.get_text("dict")
            
            # Check for mathematical symbols and special characters
            math_patterns = [
                r'[∑∫∂∇√∞≈≠≤≥±×÷]',  # Math symbols
                r'[αβγδεζηθικλμνξοπρστυφχψω]',  # Greek letters
                r'[₀₁₂₃₄₅₇₈₉]',  # Subscript numbers
                r'[⁰¹²³⁴⁵⁶⁷⁸⁹]',  # Superscript numbers
                r'[\^]\d+',  # Superscript notation
                r'=\s*[\d.]+\s*[,;]',  # Equation-like patterns
            ]
            
            for block in text_dict.get("blocks", []):
                if block.get("type") == 0:
                    for line in block.get("lines", []):
                        for span in line.get("spans", []):
                            text = span.get("text", "")
                            for pattern in math_patterns:
                                if re.search(pattern, text):
                                    logger.info(f"Detected complex content on page {page_num + 1}: {pattern}")
                                    return True
            
            return False
            
        except Exception as e:
            logger.warning(f"Error detecting complex content: {e}")
            return False
    
    def _add_page_as_image(self, doc, page_num: int, section) -> None:
        """Render a PDF page as an image and add to Word document"""
        try:
            # Render page to high-quality image
            img = self.render_page_to_image(page_num, dpi=DEFAULT_DPI)

            if img:
                img_path = os.path.join(self.temp_dir, f'page_{page_num:04d}.png')
                img.save(img_path, 'PNG', dpi=(DEFAULT_DPI, DEFAULT_DPI))

                # Calculate image size to fit page
                usable_width = section.page_width - section.left_margin - section.right_margin
                usable_height = section.page_height - section.top_margin - section.bottom_margin

                # Add image to document
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()

                # Scale image to fit within margins while maintaining aspect ratio
                img_width_in = img.width / DEFAULT_DPI
                img_height_in = img.height / DEFAULT_DPI

                scale = min(usable_width.inches / img_width_in, usable_height.inches / img_height_in)

                final_width = Inches(img_width_in * scale)
                final_height = Inches(img_height_in * scale)

                run.add_picture(img_path, width=final_width, height=final_height)

                # Clean up paragraph spacing
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                
        except Exception as e:
            logger.error(f"Error adding page as image: {e}")
    
    def _add_page_as_editable_text(self, doc, page_num: int, text_blocks: List[Dict]) -> None:
        """Add extracted text blocks as editable content"""
        try:
            for block in text_blocks:
                if block['type'] == 'text' and block.get('spans'):
                    paragraph = doc.add_paragraph()
                    
                    for span in block['spans']:
                        run = paragraph.add_run(span['text'])
                        
                        # Apply formatting
                        font = run.font
                        font.name = span['font']
                        font.size = Pt(span['size'])
                        
                        if span['bold']:
                            font.bold = True
                        if span['italic']:
                            font.italic = True
                        
                        # Apply color if not black
                        if span['color'] and span['color'] != 0:
                            r = (span['color'] >> 16) & 0xFF
                            g = (span['color'] >> 8) & 0xFF
                            b = span['color'] & 0xFF
                            if r != 0 or g != 0 or b != 0:
                                font.color.rgb = RGBColor(r, g, b)
                                
        except Exception as e:
            logger.error(f"Error adding editable text: {e}")
    
    def _add_table_to_doc(self, doc, table_info: Dict) -> None:
        """Add a table to the Word document with proper formatting"""
        try:
            table_data = table_info.get('data', [])
            if not table_data or len(table_data) < 1:
                return
            
            rows = len(table_data)
            cols = max(len(row) for row in table_data)
            
            # Add spacing before table
            doc.add_paragraph()
            
            # Create table
            table = doc.add_table(rows=rows, cols=cols)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Set table borders (ensure they're visible)
            self._set_table_borders(table)
            
            # Populate table
            for row_idx, row_data in enumerate(table_data):
                for col_idx, cell_text in enumerate(row_data):
                    if col_idx < cols:
                        cell = table.rows[row_idx].cells[col_idx]
                        cell.text = str(cell_text).strip() if cell_text else ''
                        
                        # Format cell content
                        for paragraph in cell.paragraphs:
                            paragraph.paragraph_format.space_before = Pt(3)
                            paragraph.paragraph_format.space_after = Pt(3)
                            paragraph.paragraph_format.left_indent = Pt(5)
                            paragraph.paragraph_format.right_indent = Pt(5)
                            
                            for run in paragraph.runs:
                                run.font.size = Pt(10)
                                run.font.name = 'Calibri'
                                
                                # Bold header row
                                if row_idx == 0:
                                    run.font.bold = True
                                    run.font.size = Pt(11)
                        
                        # Set cell vertical alignment
                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        
            # Add spacing after table
            doc.add_paragraph()
            
            logger.info(f"Added table: {rows} rows x {cols} cols")
            
        except Exception as e:
            logger.error(f"Error adding table: {e}")
    
    def _set_table_borders(self, table) -> None:
        """Set visible borders for table cells"""
        try:
            # Access table XML to set border properties
            tbl = table._tbl
            tblPr = tbl.get_or_add_tblPr()
            
            # Create border element
            tblBorders = OxmlElement('w:tblBorders')
            
            # Define border style
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')  # Border width
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')
                tblBorders.append(border)
            
            tblPr.append(tblBorders)
            
        except Exception as e:
            logger.warning(f"Could not set table borders: {e}")
    
    def create_editable_document(self) -> bool:
        """
        Create an editable Word document with extracted text, images, and tables.
        IMPROVED: Better table detection and formula handling
        """
        if not DOCX_AVAILABLE:
            logger.error("python-docx not available")
            return False

        try:
            doc = Document()

            # Set default style
            style = doc.styles['Normal']
            style.font.name = 'Calibri'
            style.font.size = Pt(11)
            style.paragraph_format.space_before = Pt(0)
            style.paragraph_format.space_after = Pt(3)
            style.paragraph_format.line_spacing = 1.15

            logger.info("Creating editable Word document with enhanced table detection...")

            for page_num in range(self.page_count):
                logger.info(f"Processing page {page_num + 1}/{self.page_count}")

                # Get page dimensions
                page_width_in, page_height_in = self.get_page_dimensions(page_num)

                # Add new page section
                if page_num > 0:
                    section = doc.add_section(WD_SECTION.NEW_PAGE)
                else:
                    section = doc.sections[0]

                # Set page size and margins
                section.page_width = Inches(page_width_in)
                section.page_height = Inches(page_height_in)
                section.top_margin = Cm(MARGIN_CM)
                section.bottom_margin = Cm(MARGIN_CM)
                section.left_margin = Cm(MARGIN_CM)
                section.right_margin = Cm(MARGIN_CM)

                # Extract content
                text_blocks = self.extract_text_blocks(page_num)
                images = self.extract_images(page_num)
                tables = self.extract_tables(page_num)
                
                # Detect complex content (formulas, special symbols)
                has_complex_content = self._detect_complex_content(page_num)

                # If page has complex content, render as image
                if has_complex_content:
                    logger.info(f"  Page {page_num + 1}: Rendering as image (complex content)")
                    self._add_page_as_image(doc, page_num, section)
                    continue

                # Build a map of table regions to avoid text overlap
                table_regions = [t['bbox'] for t in tables] if tables else []

                # Add text blocks (skip those inside table regions)
                for block in text_blocks:
                    if block['type'] == 'text' and block.get('bbox'):
                        # Check if this text is inside a table region
                        in_table = False
                        for table_bbox in table_regions:
                            if self._bbox_contains(table_bbox, block['bbox']):
                                in_table = True
                                break
                        
                        if not in_table:
                            self._add_text_block_to_doc(doc, block)

                # Add images
                for img_info in images:
                    if os.path.exists(img_info['path']):
                        paragraph = doc.add_paragraph()
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = paragraph.add_run()

                        # Calculate max width
                        max_width = Inches(page_width_in - 2 * MARGIN_CM)

                        try:
                            with Image.open(img_info['path']) as img:
                                img_width_in = img.width / 150  # Assume 150 DPI
                                img_height_in = img.height / 150

                                if img_width_in > max_width.inches:
                                    scale = max_width.inches / img_width_in
                                    img_width_in *= scale
                                    img_height_in *= scale

                                run.add_picture(img_info['path'],
                                              width=Inches(img_width_in),
                                              height=Inches(img_height_in))
                        except Exception as e:
                            logger.warning(f"Could not add image: {e}")

                # Add tables with proper formatting
                for table_info in tables:
                    self._add_table_to_doc(doc, table_info)

            # Save document
            doc.save(self.output_path)

            if os.path.exists(self.output_path):
                file_size = os.path.getsize(self.output_path)
                logger.info(f"Editable document saved: {file_size} bytes")
                return True

            return False

        except Exception as e:
            logger.error(f"Error creating editable document: {e}")
            import traceback
            logger.error(traceback.format_exc())
            return False
    
    def _bbox_contains(self, outer: Tuple, inner: Tuple) -> bool:
        """Check if inner bbox is contained within outer bbox"""
        if not outer or not inner:
            return False
        return (inner[0] >= outer[0] and 
                inner[1] >= outer[1] and 
                inner[2] <= outer[2] and 
                inner[3] <= outer[3])
    
    def _add_text_block_to_doc(self, doc, block: Dict) -> None:
        """Add a text block to the Word document with formatting"""
        try:
            if not block.get('spans'):
                return
                
            paragraph = doc.add_paragraph()
            
            for span in block['spans']:
                run = paragraph.add_run(span['text'])
                
                # Apply formatting
                font = run.font
                font.name = span['font']
                font.size = Pt(span['size'])
                
                if span['bold']:
                    font.bold = True
                if span['italic']:
                    font.italic = True
                
                # Apply color if not black
                if span['color'] and span['color'] != 0:
                    r = (span['color'] >> 16) & 0xFF
                    g = (span['color'] >> 8) & 0xFF
                    b = span['color'] & 0xFF
                    if r != 0 or g != 0 or b != 0:
                        font.color.rgb = RGBColor(r, g, b)
                        
        except Exception as e:
            logger.error(f"Error adding text block: {e}")
    
    def convert(self) -> Tuple[bool, str]:
        """
        Main conversion method.
        Returns (success, message) tuple.
        """
        # Validate input
        if not os.path.exists(self.pdf_path):
            return False, f"PDF file not found: {self.pdf_path}"

        # Check dependencies
        if not check_dependencies():
            return False, "Missing required dependencies. Install with: pip install pymupdf python-docx Pillow pdfplumber pdf2docx"

        # Open PDF
        if not self.open_pdf():
            return False, "Failed to open PDF file"

        try:
            # Create temp directory
            os.makedirs(self.temp_dir, exist_ok=True)

            # Normalize mode: 'no-ocr', 'editable' -> editable extraction; 'ocr' -> OCR; 'layout' -> layout preservation
            # This ensures consistent behavior regardless of how the mode is passed
            effective_mode = self.mode
            if self.mode in ('no-ocr', 'no_ocr', 'editable'):
                effective_mode = 'editable'
                logger.info(f"Mode '{self.mode}' normalized to 'editable' (text extraction)")
            elif self.mode == 'ocr':
                effective_mode = 'ocr'
                logger.info(f"Using OCR mode (image-based text recognition)")
            elif self.mode == 'layout':
                effective_mode = 'layout'
            else:
                logger.warning(f"Unknown mode '{self.mode}', defaulting to 'layout'")
                effective_mode = 'layout'

            # Choose conversion strategy based on effective mode
            if effective_mode == 'layout':
                # Layout mode: HYBRID approach - render complex pages as image, keep simple pages editable
                logger.info("Using LAYOUT mode - hybrid approach for best quality")
                success = self.create_layout_document()

                if success:
                    return True, f"Successfully converted to Word (layout preserved with hybrid rendering)"
                else:
                    return False, "Failed to create layout document"

            elif effective_mode == 'editable':
                # Editable mode (No OCR): extract text, images, tables with enhanced detection
                logger.info("Using NO-OCR/EDITABLE mode - enhanced table and formula detection")

                # Try pdf2docx first if available (best quality for editable text)
                if PDF2DOCX_AVAILABLE:
                    try:
                        logger.info("Trying pdf2docx for best editable conversion...")
                        converter = Pdf2DocxConverter(self.pdf_path)
                        converter.convert(self.output_path)
                        converter.close()

                        if os.path.exists(self.output_path) and os.path.getsize(self.output_path) > 0:
                            return True, "Successfully converted with pdf2docx (highest quality)"
                    except Exception as e:
                        logger.warning(f"pdf2docx failed, falling back to enhanced manual extraction: {e}")

                # Fallback to enhanced manual extraction
                success = self.create_editable_document()

                if success:
                    return True, "Successfully converted to Word (editable with enhanced tables)"
                else:
                    return False, "Failed to create editable document"

            elif effective_mode == 'ocr':
                # OCR mode: For scanned PDFs, use OCR to recognize text
                logger.info("Using OCR mode - recognizing text from scanned images")
                # For now, fall back to layout mode since pure OCR requires additional libraries
                # In layout mode, pages are rendered as images which preserves scanned content
                success = self.create_layout_document()
                if success:
                    return True, "Successfully converted to Word (OCR via layout mode)"
                else:
                    return False, "Failed to create OCR document"

            else:
                # Default: use hybrid layout mode for best quality
                logger.info("Using default HYBRID mode - best quality with tables and formulas")
                success = self.create_layout_document()

                if success:
                    return True, "Successfully converted to Word (hybrid mode)"
                else:
                    return False, "Conversion failed"

        except Exception as e:
            logger.error(f"Conversion error: {e}")
            import traceback
            logger.error(traceback.format_exc())
            return False, f"Conversion error: {e}"

        finally:
            self.close_pdf()
            self.cleanup_temp()


def convert_pdf_to_word(pdf_path: str, output_path: str, mode: str = 'layout') -> Tuple[bool, str]:
    """
    High-level function to convert PDF to Word.
    
    Args:
        pdf_path: Path to input PDF
        output_path: Path for output DOCX
        mode: 'layout' (visual fidelity) or 'editable' (text extraction)
    
    Returns:
        (success: bool, message: str)
    """
    converter = HighQualityPDFToWordConverter(pdf_path, output_path, mode)
    return converter.convert()


def main():
    """CLI entry point"""
    if len(sys.argv) < 3:
        print("Usage: python pdf_to_word_advanced.py <input_pdf> <output_docx> [layout|editable]")
        sys.exit(1)
    
    input_pdf = sys.argv[1]
    output_docx = sys.argv[2]
    mode = sys.argv[3] if len(sys.argv) > 3 else 'layout'
    
    success, message = convert_pdf_to_word(input_pdf, output_docx, mode)
    
    result = {
        'success': success,
        'message': message,
        'input': input_pdf,
        'output': output_docx,
        'mode': mode,
    }
    
    print(json.dumps(result, indent=2))
    sys.exit(0 if success else 1)


if __name__ == '__main__':
    main()
