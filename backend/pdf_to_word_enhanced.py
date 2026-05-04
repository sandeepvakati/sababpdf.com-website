#!/usr/bin/env python3
"""
Enhanced PDF to Word Converter - High Quality Version
Optimized for accurate text extraction, table preservation, and layout fidelity
Supports: Text, Images, Tables, Formulas, and Complex Layouts
"""

import sys
import os
from pathlib import Path
import json
import logging
import re
import shutil
import subprocess
from typing import List, Tuple, Optional, Dict, Any
from io import BytesIO

# Configure logging
LOG_LEVEL_NAME = os.environ.get('PDF_TO_WORD_LOG_LEVEL', 'ERROR').upper()
LOG_LEVEL = getattr(logging, LOG_LEVEL_NAME, logging.ERROR)
logging.basicConfig(
    level=LOG_LEVEL,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Import required packages
FITZ_AVAILABLE = False
try:
    import fitz  # PyMuPDF
    FITZ_AVAILABLE = True
except ImportError as e:
    logger.error(f"PyMuPDF import error: {e}")

PDFPLUMBER_AVAILABLE = False
try:
    import pdfplumber
    import logging
    logging.getLogger('pdfminer').setLevel(logging.ERROR)
    logging.getLogger('pdfplumber').setLevel(logging.ERROR)
    PDFPLUMBER_AVAILABLE = True
except ImportError as e:
    logger.warning(f"pdfplumber import error: {e}")

DOCX_AVAILABLE = False
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm, Emu, Twips
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError as e:
    logger.error(f"python-docx import error: {e}")

PIL_AVAILABLE = False
try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError as e:
    logger.warning(f"Pillow import error: {e}")

PDF2DOCX_AVAILABLE = False
try:
    import pdf2docx  # noqa: F401
    PDF2DOCX_AVAILABLE = True
except ImportError as e:
    logger.warning(f"pdf2docx import error: {e}")

# Constants
DEFAULT_DPI = 300
MARGIN_CM = 2.0
DEFAULT_PDF2DOCX_TIMEOUT_SECONDS = 180
DEFAULT_PDF2DOCX_MAX_PAGES = 120
DEFAULT_PDF2DOCX_MAX_FILE_SIZE_MB = 15.0
DEFAULT_RICH_EDITABLE_MAX_PAGES = 24
DEFAULT_RICH_EDITABLE_MAX_FILE_SIZE_MB = 8.0


def _read_positive_int_env(name: str, default: int) -> int:
    try:
        value = int(os.environ.get(name, default))
        return value if value > 0 else default
    except (TypeError, ValueError):
        return default


def _read_positive_float_env(name: str, default: float) -> float:
    try:
        value = float(os.environ.get(name, default))
        return value if value > 0 else default
    except (TypeError, ValueError):
        return default


PDF2DOCX_TIMEOUT_SECONDS = _read_positive_int_env(
    'PDF_TO_WORD_PDF2DOCX_TIMEOUT_SECONDS',
    DEFAULT_PDF2DOCX_TIMEOUT_SECONDS,
)
PDF2DOCX_MAX_PAGES = _read_positive_int_env(
    'PDF_TO_WORD_PDF2DOCX_MAX_PAGES',
    DEFAULT_PDF2DOCX_MAX_PAGES,
)
PDF2DOCX_MAX_FILE_SIZE_MB = _read_positive_float_env(
    'PDF_TO_WORD_PDF2DOCX_MAX_FILE_SIZE_MB',
    DEFAULT_PDF2DOCX_MAX_FILE_SIZE_MB,
)
RICH_EDITABLE_MAX_PAGES = _read_positive_int_env(
    'PDF_TO_WORD_RICH_EDITABLE_MAX_PAGES',
    DEFAULT_RICH_EDITABLE_MAX_PAGES,
)
RICH_EDITABLE_MAX_FILE_SIZE_MB = _read_positive_float_env(
    'PDF_TO_WORD_RICH_EDITABLE_MAX_FILE_SIZE_MB',
    DEFAULT_RICH_EDITABLE_MAX_FILE_SIZE_MB,
)


def check_dependencies():
    """Check and report available dependencies"""
    deps = {
        'PyMuPDF': FITZ_AVAILABLE,
        'python-docx': DOCX_AVAILABLE,
        'Pillow': PIL_AVAILABLE,
    }

    missing = [name for name, available in deps.items() if not available]

    if missing:
        logger.error(f"Missing required dependencies: {', '.join(missing)}")
        logger.error("Install with: pip install pymupdf python-docx Pillow")
        return False

    return True


class EnhancedPDFToWordConverter:
    """
    Enhanced PDF to Word converter with improved quality.
    Strategy: Try pdf2docx first (best quality), fallback to custom extraction.
    """

    def __init__(self, pdf_path: str, output_path: str, mode: str = 'layout'):
        self.pdf_path = pdf_path
        self.output_path = output_path
        self.mode = mode
        self.temp_dir = os.path.join(os.path.dirname(output_path), 'temp_conversion')
        self.pdf_doc = None
        self.page_count = 0
        self.pdfplumber_doc = None
        self.use_rich_editable_extraction = True

    def open_pdf(self) -> bool:
        """Open PDF document with PyMuPDF"""
        if not FITZ_AVAILABLE:
            logger.error("PyMuPDF is required")
            return False

        try:
            self.pdf_doc = fitz.open(self.pdf_path)
            # Detect password-protected / encrypted PDFs
            if self.pdf_doc.is_encrypted:
                self.pdf_doc.close()
                self.pdf_doc = None
                logger.error("PDF is password-protected or encrypted")
                return False
            self.page_count = len(self.pdf_doc)
            logger.info(f"Opened PDF: {self.page_count} pages")
            return True
        except Exception as e:
            err_str = str(e).lower()
            if 'encrypted' in err_str or 'password' in err_str or 'closed' in err_str:
                logger.error(f"PDF is password-protected or encrypted: {e}")
            else:
                logger.error(f"Failed to open PDF: {e}")
            return False

    def open_pdfplumber(self) -> bool:
        """Open PDF with pdfplumber for better table detection"""
        if not PDFPLUMBER_AVAILABLE:
            return False

        try:
            self.pdfplumber_doc = pdfplumber.open(self.pdf_path)
            logger.info("Opened pdfplumber document")
            return True
        except Exception as e:
            logger.warning(f"Failed to open pdfplumber: {e}")
            return False

    def close_pdf(self):
        """Close PDF documents"""
        if self.pdf_doc:
            self.pdf_doc.close()
            self.pdf_doc = None
        if self.pdfplumber_doc:
            self.pdfplumber_doc.close()
            self.pdfplumber_doc = None

    def _get_file_size_mb(self) -> float:
        try:
            return os.path.getsize(self.pdf_path) / (1024 * 1024)
        except OSError:
            return 0.0

    def cleanup_temp(self):
        """Clean up temporary files"""
        if os.path.exists(self.temp_dir):
            try:
                shutil.rmtree(self.temp_dir, ignore_errors=True)
            except Exception as e:
                logger.warning(f"Could not clean temp dir: {e}")

    def get_page_dimensions(self, page_num: int) -> Tuple[float, float]:
        """Get page dimensions in inches"""
        page = self.pdf_doc[page_num]
        rect = page.rect
        return rect.width / 72, rect.height / 72

    def render_page_to_image(self, page_num: int, dpi: int = DEFAULT_DPI) -> Optional[Image.Image]:
        """Render a PDF page to a high-quality PIL Image"""
        if not FITZ_AVAILABLE or not PIL_AVAILABLE:
            return None

        try:
            page = self.pdf_doc[page_num]
            zoom = dpi / 72
            mat = fitz.Matrix(zoom, zoom)

            pix = page.get_pixmap(matrix=mat, alpha=False)
            img_data = pix.tobytes('png')
            img = Image.open(BytesIO(img_data))

            return img.convert('RGB')
        except Exception as e:
            logger.error(f"Failed to render page {page_num}: {e}")
            return None

    def extract_text_with_positions(self, page_num: int) -> List[Dict]:
        """
        Extract text blocks with position and formatting information.
        Improved: Better handling of text fragments and spacing.
        """
        if not FITZ_AVAILABLE:
            return []

        try:
            page = self.pdf_doc[page_num]
            # Use 'dict' mode with whitespace preservation
            text_dict = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE | fitz.TEXT_MEDIABOX_CLIP)

            blocks = []
            for block in text_dict.get("blocks", []):
                if block.get("type") == 0:  # Text block
                    for line in block.get("lines", []):
                        spans = []
                        for span in line.get("spans", []):
                            text = span.get("text", "").strip()
                            if text:
                                spans.append({
                                    'text': text,
                                    'font': span.get("font", "Arial"),
                                    'size': span.get("size", 11),
                                    'color': span.get("color", 0),
                                    'bold': 'bold' in span.get("font", "").lower(),
                                    'italic': 'italic' in span.get("font", "").lower() or
                                             'oblique' in span.get("font", "").lower(),
                                })

                        if spans:
                            blocks.append({
                                'type': 'text',
                                'spans': spans,
                                'bbox': line.get("bbox"),
                                'y_position': line.get("bbox", [0, 0, 0, 0])[1],
                            })

            # Sort blocks by y-position (top to bottom)
            blocks.sort(key=lambda b: b['y_position'])

            return blocks
        except Exception as e:
            logger.error(f"Error extracting text from page {page_num}: {e}")
            return []

    def extract_images(self, page_num: int) -> List[Dict]:
        """Extract images from a PDF page"""
        if not FITZ_AVAILABLE:
            return []

        images = []
        os.makedirs(self.temp_dir, exist_ok=True)

        try:
            page = self.pdf_doc[page_num]
            image_list = page.get_images(full=True)

            for img_index, img_info in enumerate(image_list):
                try:
                    xref = img_info[0]
                    base_image = fitz.Pixmap(self.pdf_doc, xref)

                    if base_image.n - base_image.alpha > 3:  # CMYK
                        base_image = fitz.Pixmap(fitz.csRGB, base_image)

                    img_path = os.path.join(self.temp_dir, f'p{page_num}_img{img_index}.png')
                    base_image.save(img_path)

                    img_rect = img_info[1:5] if len(img_info) > 4 else None

                    images.append({
                        'type': 'image',
                        'path': img_path,
                        'bbox': img_rect,
                        'index': img_index,
                    })

                    base_image = None  # Free memory
                except Exception as e:
                    logger.warning(f"Could not extract image {img_index}: {e}")

            logger.info(f"Extracted {len(images)} images from page {page_num}")

        except Exception as e:
            logger.error(f"Error extracting images from page {page_num}: {e}")

        return images

    def extract_tables_with_pdfplumber(self, page_num: int) -> List[Dict]:
        """
        Extract tables using pdfplumber for better accuracy.
        Returns list of tables with cell data.
        """
        if not PDFPLUMBER_AVAILABLE or not self.pdfplumber_doc:
            return []

        tables = []

        try:
            if page_num < len(self.pdfplumber_doc.pages):
                pdf_page = self.pdfplumber_doc.pages[page_num]
                pdf_tables = pdf_page.find_tables()

                for table_idx, table in enumerate(pdf_tables):
                    table_data = table.extract()
                    if table_data and len(table_data) > 0:
                        # Clean and validate table data
                        max_cols = max(len(row) for row in table_data)
                        cleaned_data = []
                        for row in table_data:
                            cleaned_row = [str(cell).strip() if cell is not None else '' for cell in row]
                            while len(cleaned_row) < max_cols:
                                cleaned_row.append('')
                            cleaned_data.append(cleaned_row[:max_cols])

                        tables.append({
                            'type': 'table',
                            'data': cleaned_data,
                            'bbox': table.bbox,
                            'index': table_idx,
                        })

                logger.info(f"Extracted {len(tables)} tables from page {page_num} via pdfplumber")

        except Exception as e:
            logger.warning(f"Error extracting tables with pdfplumber: {e}")

        return tables

    def _detect_complex_content(
        self,
        page_num: int,
        text_blocks: Optional[List[Dict]] = None,
        tables: Optional[List[Dict]] = None,
        images: Optional[List[Dict]] = None,
    ) -> bool:
        """Detect when a page should be preserved visually instead of rebuilt as editable text."""
        if not FITZ_AVAILABLE:
            return False

        try:
            page = self.pdf_doc[page_num]
            text_dict = page.get_text("dict")
            page_area = max(page.rect.width * page.rect.height, 1)

            if tables and len(tables) > 0:
                return True

            if images and len(images) >= 2:
                return True

            drawings = page.get_drawings()
            filled_regions = 0
            for drawing in drawings:
                fill = drawing.get("fill")
                fill_opacity = drawing.get("fill_opacity", 1 if fill else 0)
                rect = drawing.get("rect")

                if fill and fill_opacity and rect:
                    filled_regions += 1
                    area_ratio = (rect.width * rect.height) / page_area
                    if area_ratio >= 0.12:
                        return True

            if len(drawings) >= 24 or filled_regions >= 6:
                return True

            candidate_blocks = text_blocks or []
            if candidate_blocks:
                block_starts = sorted({
                    int((block.get('bbox', [0, 0, 0, 0])[0] // 24) * 24)
                    for block in candidate_blocks
                    if block.get('bbox')
                })
                if len(block_starts) >= 2 and (block_starts[-1] - block_starts[0]) >= 140:
                    return True

            # Patterns for complex content
            math_patterns = [
                r'[∑∫∂√∞≈≠≤≥±×÷]',  # Math symbols
                r'[αβγδεζηθικλμνξοπρστυφχψω]',  # Greek letters
                r'[₀₁₂₃₄₅₆₇₈₉]',  # Subscript numbers
                r'[⁰¹²³⁴⁵⁶⁸⁹]',  # Superscript numbers
            ]

            total_spans = 0
            colored_spans = 0

            for block in text_dict.get("blocks", []):
                if block.get("type") == 0:
                    for line in block.get("lines", []):
                        for span in line.get("spans", []):
                            text = span.get("text", "")
                            total_spans += 1
                            color = span.get("color", 0)
                            if color not in (0, None):
                                colored_spans += 1
                            for pattern in math_patterns:
                                if re.search(pattern, text):
                                    return True

            if total_spans >= 12 and (colored_spans / total_spans) >= 0.18:
                return True

            return False

        except Exception as e:
            logger.warning(f"Error detecting complex content: {e}")
            return False

    def _add_table_to_doc(self, doc, table_info: Dict) -> None:
        """Add a table to the Word document with proper formatting"""
        try:
            table_data = table_info.get('data', [])
            if not table_data or len(table_data) < 1:
                return

            rows = len(table_data)
            cols = max(len(row) for row in table_data)

            # Add spacing before table
            doc.add_paragraph()

            # Create table
            table = doc.add_table(rows=rows, cols=cols)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Set table borders
            self._set_table_borders(table)

            # Populate table
            for row_idx, row_data in enumerate(table_data):
                for col_idx, cell_text in enumerate(row_data):
                    if col_idx < cols:
                        cell = table.rows[row_idx].cells[col_idx]
                        cell.text = str(cell_text).strip() if cell_text else ''

                        # Format cell content
                        for paragraph in cell.paragraphs:
                            paragraph.paragraph_format.space_before = Pt(2)
                            paragraph.paragraph_format.space_after = Pt(2)

                            for run in paragraph.runs:
                                run.font.size = Pt(10)
                                run.font.name = 'Calibri'

                                # Bold header row
                                if row_idx == 0:
                                    run.font.bold = True
                                    run.font.size = Pt(11)

                        # Set cell vertical alignment
                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            # Add spacing after table
            doc.add_paragraph()

            logger.info(f"Added table: {rows} rows x {cols} cols")

        except Exception as e:
            logger.error(f"Error adding table: {e}")

    def _set_table_borders(self, table) -> None:
        """Set visible borders for table cells"""
        try:
            tbl = table._tbl
            tblPr = tbl.get_or_add_tblPr()

            tblBorders = OxmlElement('w:tblBorders')

            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')
                tblBorders.append(border)

            tblPr.append(tblBorders)

        except Exception as e:
            logger.warning(f"Could not set table borders: {e}")

    def _add_text_block_to_doc(self, doc, block: Dict) -> None:
        """Add a text block to the Word document with formatting"""
        try:
            if not block.get('spans'):
                return

            paragraph = doc.add_paragraph()

            for span in block['spans']:
                run = paragraph.add_run(span['text'])

                # Apply formatting
                font = run.font
                font.name = span['font']
                font.size = Pt(span['size'])

                if span['bold']:
                    font.bold = True
                if span['italic']:
                    font.italic = True

                # Apply color if not black
                if span['color'] and span['color'] != 0:
                    r = (span['color'] >> 16) & 0xFF
                    g = (span['color'] >> 8) & 0xFF
                    b = span['color'] & 0xFF
                    if r != 0 or g != 0 or b != 0:
                        font.color.rgb = RGBColor(r, g, b)

        except Exception as e:
            logger.error(f"Error adding text block: {e}")

    def create_editable_document(self) -> bool:
        """
        Create an editable Word document with extracted content.
        Improved: Better text extraction and table handling.
        """
        if not DOCX_AVAILABLE:
            logger.error("python-docx not available")
            return False

        try:
            doc = Document()

            # Set default style
            style = doc.styles['Normal']
            style.font.name = 'Calibri'
            style.font.size = Pt(11)
            style.paragraph_format.space_before = Pt(0)
            style.paragraph_format.space_after = Pt(3)
            style.paragraph_format.line_spacing = 1.15

            logger.info("Creating editable Word document...")

            for page_num in range(self.page_count):
                logger.info(f"Processing page {page_num + 1}/{self.page_count}")

                # Get page dimensions
                page_width_in, page_height_in = self.get_page_dimensions(page_num)

                # Add new page section
                if page_num > 0:
                    section = doc.add_section(WD_SECTION.NEW_PAGE)
                else:
                    section = doc.sections[0]

                # Set page size and margins
                section.page_width = Inches(page_width_in)
                section.page_height = Inches(page_height_in)
                section.top_margin = Cm(MARGIN_CM)
                section.bottom_margin = Cm(MARGIN_CM)
                section.left_margin = Cm(MARGIN_CM)
                section.right_margin = Cm(MARGIN_CM)

                # Extract content
                text_blocks = self.extract_text_with_positions(page_num)
                if self.use_rich_editable_extraction:
                    images = self.extract_images(page_num)
                    tables = self.extract_tables_with_pdfplumber(page_num)
                else:
                    images = []
                    tables = []

                # Keep visually complex pages intact instead of rebuilding them poorly as editable text.
                has_complex_content = self._detect_complex_content(
                    page_num,
                    text_blocks=text_blocks,
                    tables=tables,
                    images=images,
                )

                # If page has complex content, render as image
                if has_complex_content:
                    logger.info(f"  Page {page_num + 1}: Rendering as image (complex content)")
                    self._add_page_as_image(doc, page_num, section)
                    continue

                # Build table regions to avoid text overlap
                table_regions = [t['bbox'] for t in tables] if tables else []

                # Add text blocks (skip those inside table regions)
                for block in text_blocks:
                    if block['type'] == 'text' and block.get('bbox'):
                        in_table = False
                        for table_bbox in table_regions:
                            if self._bbox_contains(table_bbox, block['bbox']):
                                in_table = True
                                break

                        if not in_table:
                            self._add_text_block_to_doc(doc, block)

                # Add images
                for img_info in images:
                    if os.path.exists(img_info['path']):
                        paragraph = doc.add_paragraph()
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = paragraph.add_run()

                        max_width = Inches(page_width_in - 2 * MARGIN_CM)

                        try:
                            with Image.open(img_info['path']) as img:
                                img_width_in = img.width / 150
                                img_height_in = img.height / 150

                                if img_width_in > max_width.inches:
                                    scale = max_width.inches / img_width_in
                                    img_width_in *= scale
                                    img_height_in *= scale

                                run.add_picture(img_info['path'],
                                              width=Inches(img_width_in),
                                              height=Inches(img_height_in))
                        except Exception as e:
                            logger.warning(f"Could not add image: {e}")

                # Add tables
                for table_info in tables:
                    self._add_table_to_doc(doc, table_info)

            # Save document
            doc.save(self.output_path)

            if os.path.exists(self.output_path):
                file_size = os.path.getsize(self.output_path)
                logger.info(f"Editable document saved: {file_size} bytes")
                return True

            return False

        except Exception as e:
            logger.error(f"Error creating editable document: {e}")
            import traceback
            logger.error(traceback.format_exc())
            return False

    def _add_page_as_image(self, doc, page_num: int, section) -> None:
        """Render a PDF page as an image and add to Word document"""
        try:
            img = self.render_page_to_image(page_num, dpi=DEFAULT_DPI)

            if img:
                img_path = os.path.join(self.temp_dir, f'page_{page_num:04d}.png')
                img.save(img_path, 'PNG', dpi=(DEFAULT_DPI, DEFAULT_DPI))

                # Get usable dimensions
                usable_width_in = float(section.page_width.inches) - float(section.left_margin.inches) - float(section.right_margin.inches)
                usable_height_in = float(section.page_height.inches) - float(section.top_margin.inches) - float(section.bottom_margin.inches)

                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()

                img_width_in = img.width / DEFAULT_DPI
                img_height_in = img.height / DEFAULT_DPI

                scale = min(usable_width_in / img_width_in, usable_height_in / img_height_in)

                final_width = Inches(img_width_in * scale)
                final_height = Inches(img_height_in * scale)

                run.add_picture(img_path, width=final_width, height=final_height)

                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0

        except Exception as e:
            logger.error(f"Error adding page as image: {e}")

    def _bbox_contains(self, outer: Tuple, inner: Tuple) -> bool:
        """Check if inner bbox is contained within outer bbox"""
        if not outer or not inner:
            return False
        return (inner[0] >= outer[0] and
                inner[1] >= outer[1] and
                inner[2] <= outer[2] and
                inner[3] <= outer[3])

    def _should_try_pdf2docx(self, effective_mode: str) -> Tuple[bool, str]:
        """Use pdf2docx only for smaller editable PDFs where it is most reliable."""
        if effective_mode != 'editable':
            return False, f"Skipping pdf2docx for mode '{effective_mode}'."

        if not PDF2DOCX_AVAILABLE:
            return False, 'pdf2docx is not installed.'

        if self.page_count > PDF2DOCX_MAX_PAGES:
            return False, f"Skipping pdf2docx because the PDF has {self.page_count} pages."

        file_size_mb = self._get_file_size_mb()

        if file_size_mb > PDF2DOCX_MAX_FILE_SIZE_MB:
            return False, f"Skipping pdf2docx because the PDF is {file_size_mb:.1f} MB."

        return True, 'Trying pdf2docx first for a smaller editable PDF.'

    def _should_use_rich_editable_extraction(self) -> Tuple[bool, str]:
        """
        Rich editable extraction preserves more tables and images, but it is slower.
        For larger or more complex PDFs we switch to a faster text-first DOCX path.
        """
        file_size_mb = self._get_file_size_mb()

        if self.page_count > RICH_EDITABLE_MAX_PAGES:
            return False, (
                f"Using fast editable mode because the PDF has {self.page_count} pages. "
                "Skipping expensive table and embedded-image extraction."
            )

        if file_size_mb > RICH_EDITABLE_MAX_FILE_SIZE_MB:
            return False, (
                f"Using fast editable mode because the PDF is {file_size_mb:.1f} MB. "
                "Skipping expensive table and embedded-image extraction."
            )

        return True, 'Using rich editable mode with table and image extraction.'

    def _run_pdf2docx_with_timeout(self) -> Tuple[bool, str]:
        """Run pdf2docx in an isolated subprocess so a hung engine can time out cleanly."""
        intermediate_output_path = os.path.join(self.temp_dir, 'pdf2docx-output.docx')
        script = """
import sys
from pdf2docx import Converter

input_path = sys.argv[1]
output_path = sys.argv[2]
converter = Converter(input_path)
try:
    converter.convert(output_path)
finally:
    converter.close()
"""

        try:
            result = subprocess.run(
                [sys.executable, '-c', script, self.pdf_path, intermediate_output_path],
                capture_output=True,
                text=True,
                timeout=PDF2DOCX_TIMEOUT_SECONDS,
                check=False,
            )
        except subprocess.TimeoutExpired:
            return False, f'pdf2docx timed out after {PDF2DOCX_TIMEOUT_SECONDS} seconds.'
        except Exception as exc:
            return False, f'pdf2docx could not start: {exc}'

        if result.returncode != 0:
            details = (result.stderr or result.stdout or '').strip()
            return False, details or f'pdf2docx exited with code {result.returncode}.'

        if not os.path.exists(intermediate_output_path) or os.path.getsize(intermediate_output_path) == 0:
            return False, 'pdf2docx finished without creating a DOCX file.'

        shutil.move(intermediate_output_path, self.output_path)
        return True, 'Successfully converted with pdf2docx (highest quality)'

    def convert(self) -> Tuple[bool, str]:
        """
        Main conversion method with quality-first approach.
        Strategy:
        1. Try pdf2docx (best quality for editable text)
        2. Fallback to custom extraction with enhanced table detection
        """
        # Validate input
        if not os.path.exists(self.pdf_path):
            return False, f"PDF file not found: {self.pdf_path}"

        # Check dependencies
        if not check_dependencies():
            return False, "Missing required dependencies"

        # Open PDF
        if not self.open_pdf():
            # Check if it's an encryption issue
            import sys as _sys
            if FITZ_AVAILABLE:
                try:
                    _test = fitz.open(self.pdf_path)
                    if _test.is_encrypted:
                        _test.close()
                        return False, "This PDF is password-protected or encrypted. Please remove the password first using the Unlock PDF tool."
                    _test.close()
                except Exception as _e:
                    _es = str(_e).lower()
                    if 'encrypted' in _es or 'password' in _es or 'closed' in _es:
                        return False, "This PDF is password-protected or encrypted. Please remove the password first using the Unlock PDF tool."
            return False, "Failed to open the PDF file. The file may be corrupted, password-protected, or not a valid PDF."

        try:
            # Create temp directory
            os.makedirs(self.temp_dir, exist_ok=True)

            # Normalize mode
            effective_mode = self.mode
            if self.mode in ('no-ocr', 'no_ocr', 'editable'):
                effective_mode = 'editable'
            elif self.mode == 'ocr':
                effective_mode = 'ocr'
            elif self.mode == 'layout':
                effective_mode = 'layout'
            else:
                effective_mode = 'layout'

            if effective_mode == 'editable':
                self.use_rich_editable_extraction, editable_reason = self._should_use_rich_editable_extraction()
                logger.info(editable_reason)
                if self.use_rich_editable_extraction:
                    self.open_pdfplumber()
            else:
                self.use_rich_editable_extraction = False

            should_try_pdf2docx, pdf2docx_reason = self._should_try_pdf2docx(effective_mode)
            logger.info(pdf2docx_reason)

            # STRATEGY 1: Try pdf2docx first only for smaller editable PDFs
            if should_try_pdf2docx:
                success, message = self._run_pdf2docx_with_timeout()
                if success:
                    return True, message
                logger.warning(f"pdf2docx fallback triggered: {message}")

            # STRATEGY 2: Custom extraction
            if effective_mode == 'editable':
                logger.info("Using custom extraction for editable document...")
                success = self.create_editable_document()
                if success:
                    return True, "Successfully converted to Word (editable)"
                else:
                    return False, "Failed to create editable document"

            elif effective_mode == 'layout':
                logger.info("Using layout mode (rendering pages as images)...")
                success = self.create_layout_document()
                if success:
                    return True, "Successfully converted to Word (layout preserved)"
                else:
                    return False, "Failed to create layout document"

            elif effective_mode == 'ocr':
                logger.info("Using OCR mode (layout-based)...")
                success = self.create_layout_document()
                if success:
                    return True, "Successfully converted to Word (OCR mode)"
                else:
                    return False, "Failed to create OCR document"

            return False, "Conversion failed"

        except Exception as e:
            logger.error(f"Conversion error: {e}")
            import traceback
            logger.error(traceback.format_exc())
            err_str = str(e).lower()
            if 'encrypted' in err_str or 'password' in err_str or 'closed' in err_str:
                return False, "This PDF is password-protected or encrypted. Please remove the password first using the Unlock PDF tool."
            return False, f"Conversion error: {e}"

        finally:
            self.close_pdf()
            self.cleanup_temp()

    def create_layout_document(self) -> bool:
        """Create a Word document with layout preservation (pages as images)"""
        if not DOCX_AVAILABLE:
            logger.error("python-docx not available")
            return False

        try:
            doc = Document()

            style = doc.styles['Normal']
            style.font.name = 'Calibri'
            style.font.size = Pt(11)

            logger.info("Creating layout-preserving Word document...")

            for page_num in range(self.page_count):
                logger.info(f"Processing page {page_num + 1}/{self.page_count}")

                page_width_in, page_height_in = self.get_page_dimensions(page_num)

                if page_num > 0:
                    section = doc.add_section(WD_SECTION.NEW_PAGE)
                else:
                    section = doc.sections[0]

                section.page_width = Inches(page_width_in)
                section.page_height = Inches(page_height_in)
                section.top_margin = Cm(MARGIN_CM)
                section.bottom_margin = Cm(MARGIN_CM)
                section.left_margin = Cm(MARGIN_CM)
                section.right_margin = Cm(MARGIN_CM)

                # Render page as image
                self._add_page_as_image(doc, page_num, section)

            doc.save(self.output_path)

            if os.path.exists(self.output_path):
                file_size = os.path.getsize(self.output_path)
                logger.info(f"Layout document saved: {file_size} bytes")
                return True

            return False

        except Exception as e:
            logger.error(f"Error creating layout document: {e}")
            import traceback
            logger.error(traceback.format_exc())
            return False


def convert_pdf_to_word(pdf_path: str, output_path: str, mode: str = 'layout') -> Tuple[bool, str]:
    """
    High-level function to convert PDF to Word.

    Args:
        pdf_path: Path to input PDF
        output_path: Path for output DOCX
        mode: 'layout', 'editable', or 'ocr'

    Returns:
        (success: bool, message: str)
    """
    converter = EnhancedPDFToWordConverter(pdf_path, output_path, mode)
    return converter.convert()


def main():
    """CLI entry point"""
    if len(sys.argv) < 3:
        print("Usage: python pdf_to_word_enhanced.py <input_pdf> <output_docx> [layout|editable|ocr]")
        sys.exit(1)

    input_pdf = sys.argv[1]
    output_docx = sys.argv[2]
    mode = sys.argv[3] if len(sys.argv) > 3 else 'layout'

    success, message = convert_pdf_to_word(input_pdf, output_docx, mode)

    result = {
        'success': success,
        'message': message,
        'input': input_pdf,
        'output': output_docx,
        'mode': mode,
    }

    print(json.dumps(result, indent=2))
    sys.exit(0 if success else 1)


if __name__ == '__main__':
    main()
