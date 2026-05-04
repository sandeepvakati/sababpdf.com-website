#!/usr/bin/env python3
"""
Advanced PDF to Word Converter
Supports: Text, Images, Tables, and Mathematical Equations
Author: SababPDF
"""

import sys
import os
from pathlib import Path
import json
import logging
import re
import shutil
from typing import List, Tuple, Optional, Dict

# Configure logging
LOG_LEVEL_NAME = os.environ.get('PDF_TO_WORD_LOG_LEVEL', 'WARNING').upper()
LOG_LEVEL = getattr(logging, LOG_LEVEL_NAME, logging.WARNING)
logging.basicConfig(
    level=LOG_LEVEL,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

FAST_MODE_PAGE_THRESHOLD = 400
DEFAULT_CONVERSION_MODE = 'no-ocr'
OCR_RENDER_SCALE = 2.2

def try_import(package_name, import_name=None):
    """Try to import a package, return None if not available"""
    try:
        if import_name:
            return __import__(import_name)
        return __import__(package_name)
    except ImportError:
        return None

# Try to import required packages
fitz = try_import('pymupdf', 'fitz')
pdfplumber = try_import('pdfplumber')

# Special handling for PIL.Image
try:
    from PIL import Image
except ImportError:
    Image = None

# Other optional imports
cv2 = try_import('cv2')
numpy = try_import('numpy', 'numpy')
pytesseract = try_import('pytesseract')
pdf2image = try_import('pdf2image')
pdf2docx = try_import('pdf2docx')

try:
    from pdf2docx import Converter as Pdf2DocxConverter
except ImportError:
    Pdf2DocxConverter = None

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml
import io

MATH_SYMBOL_PATTERN = re.compile(r"[=+\-*/^√∑∫πθλμσΔΩ≤≥≈≠∞→↔±×÷∂∇∈∉⊂⊃∪∩₀-₉⁰-⁹]")
MATH_KEYWORDS = ('sin', 'cos', 'tan', 'log', 'ln', 'lim', 'max', 'min', 'mod')

def looks_like_formula(text: str) -> bool:
    """Heuristic for math-heavy lines that are safer to preserve as images."""
    if not text:
        return False

    normalized = " ".join(text.split()).strip()
    if len(normalized) < 3:
        return False

    symbol_hits = len(MATH_SYMBOL_PATTERN.findall(normalized))
    digit_hits = sum(1 for char in normalized if char.isdigit())
    alpha_hits = sum(1 for char in normalized if char.isalpha())
    keyword_hits = any(keyword in normalized.lower() for keyword in MATH_KEYWORDS)

    if symbol_hits >= 2 and (digit_hits > 0 or alpha_hits > 0 or keyword_hits):
        return True

    if any(char in normalized for char in ('∑', '∫', '√', '≤', '≥', '≈', '∞')):
        return True

    return False

def save_formula_snippet(pdf_doc, page, bbox, output_dir: str, page_num: int, block_index: int) -> Optional[str]:
    """Render a formula-like region as an image so Word keeps the visual layout."""
    if not fitz or not bbox:
        return None

    try:
        os.makedirs(output_dir, exist_ok=True)
        rect = fitz.Rect(bbox)
        padding_x = 6
        padding_y = 4
        clip = fitz.Rect(
            max(0, rect.x0 - padding_x),
            max(0, rect.y0 - padding_y),
            rect.x1 + padding_x,
            rect.y1 + padding_y,
        )

        pix = page.get_pixmap(matrix=fitz.Matrix(2.2, 2.2), clip=clip, alpha=False)
        image_path = os.path.join(output_dir, f'formula_p{page_num}_b{block_index}.png')
        pix.save(image_path)
        return image_path if os.path.exists(image_path) else None
    except Exception as e:
        logger.warning(f"Could not render formula snippet on page {page_num}: {e}")
        return None

def extract_images_from_pdf(pdf_path: str, output_dir: str) -> dict:
    """Extract all images from PDF using PyMuPDF with enhanced error handling"""
    images = {}
    
    if not fitz:
        logger.warning("PyMuPDF not available - skipping image extraction")
        return images
    
    if not os.path.exists(pdf_path):
        logger.error(f"PDF file not found: {pdf_path}")
        return images
    
    try:
        os.makedirs(output_dir, exist_ok=True)
        
        pdf_doc = fitz.open(pdf_path)
        image_count = 0
        
        logger.info(f"Starting image extraction from PDF: {pdf_path}")
        
        for page_num, page in enumerate(pdf_doc):
            try:
                image_list = page.get_images()
                
                if not image_list:
                    logger.debug(f"No images found on page {page_num}")
                    continue
                
                for img_index, img in enumerate(image_list):
                    try:
                        xref = img[0]
                        pix = fitz.Pixmap(pdf_doc, xref)
                        
                        if not os.path.exists(output_dir):
                            os.makedirs(output_dir, exist_ok=True)
                        
                        if pix.n - pix.alpha < 4:  # GRAY or RGB
                            img_path = os.path.join(
                                output_dir, 
                                f'image_p{page_num}_i{img_index}.png'
                            )
                            pix.save(img_path)
                        else:  # CMYK
                            pix = fitz.Pixmap(fitz.csRGB, pix)
                            img_path = os.path.join(
                                output_dir,
                                f'image_p{page_num}_i{img_index}.png'
                            )
                            pix.save(img_path)
                        
                        if os.path.exists(img_path):
                            images[f"p{page_num}_i{img_index}"] = img_path
                            image_count += 1
                            logger.info(f"✓ Extracted image: {img_path}")
                        else:
                            logger.warning(f"Image file not created: {img_path}")
                    
                    except Exception as img_err:
                        logger.warning(f"Could not extract image {img_index} from page {page_num}: {img_err}")
                        continue
            
            except Exception as page_err:
                logger.warning(f"Error processing page {page_num}: {page_err}")
                continue
        
        pdf_doc.close()
        logger.info(f"✓ Total images extracted: {image_count}")
        
    except Exception as e:
        logger.error(f"Fatal error extracting images: {e}")
        import traceback
        logger.error(traceback.format_exc())
    
    return images

# ─────────────────────────────────────────────────────────────────────────────
# TABLE EXTRACTION  (now returns bbox per table so text extraction can skip
# content that lives inside a table cell)
# ─────────────────────────────────────────────────────────────────────────────

def extract_tables_from_pdf(pdf_path: str) -> List[dict]:
    """Extract tables from PDF using pdfplumber.

    Each entry now contains an optional ``'bbox'`` key with the table
    bounding box ``(x0, top, x1, bottom)`` in PDF user-space coordinates.
    The caller uses these bboxes to filter out text blocks that overlap
    with a table region so that table content is never rendered *twice*
    (once as flat text and once as a proper table).
    """
    tables = []
    
    if not pdfplumber:
        logger.warning("pdfplumber not available - skipping table extraction")
        return tables
    
    if not os.path.exists(pdf_path):
        logger.error(f"PDF file not found: {pdf_path}")
        return tables
    
    try:
        logger.info("Starting table extraction")
        
        with pdfplumber.open(pdf_path) as pdf:
            total_pages = len(pdf.pages)
            logger.info(f"Processing {total_pages} pages for tables")
            
            for page_num, page in enumerate(pdf.pages):
                try:
                    # Use find_tables() so we get bounding-box information
                    # alongside the extracted data.
                    page_table_objects = page.find_tables()

                    for table_idx, table_obj in enumerate(page_table_objects):
                        try:
                            table_data = table_obj.extract()
                            if table_data and len(table_data) > 0:
                                # bbox: (x0, top, x1, bottom) in PDF coordinates
                                bbox = table_obj.bbox
                                tables.append({
                                    'page': page_num,
                                    'data': table_data,
                                    'bbox': bbox,  # NEW — used for text-block filtering
                                })
                                logger.info(
                                    f"✓ Extracted table {table_idx} from page {page_num} "
                                    f"({len(table_data)} rows, bbox={bbox})"
                                )
                        except Exception as table_err:
                            logger.warning(
                                f"Error processing table {table_idx} on page {page_num}: {table_err}"
                            )
                            continue
                
                except Exception as page_err:
                    logger.warning(f"Error extracting tables from page {page_num}: {page_err}")
                    continue
        
        logger.info(f"✓ Total tables extracted: {len(tables)}")
    
    except Exception as e:
        logger.error(f"Fatal error extracting tables: {e}")
        import traceback
        logger.error(traceback.format_exc())
    
    return tables


def _build_table_bbox_index(tables: List[dict]) -> Dict[int, List[tuple]]:
    """Return a mapping of page_num → list of (x0, top, x1, bottom) bboxes."""
    index: Dict[int, List[tuple]] = {}
    for t in tables:
        pg = t.get('page', -1)
        bbox = t.get('bbox')
        if bbox is not None:
            index.setdefault(pg, []).append(bbox)
    return index


def _block_is_inside_table(block_bbox: tuple, table_bboxes: List[tuple],
                            overlap_threshold: float = 0.45) -> bool:
    """Return True when *block_bbox* substantially overlaps any table bbox.

    Both bbox formats use (x0, y0, x1, y1) with y increasing downwards
    (PDF user-space / PyMuPDF convention).  pdfplumber also uses the same
    origin so the coordinates are directly comparable.

    overlap_threshold: fraction of the block area that must be covered by a
    table region before we decide the block is a table cell.  0.45 is a
    conservative value that catches cells reliably without false-positives on
    text that merely touches a table border.
    """
    if not table_bboxes:
        return False

    bx0, by0, bx1, by1 = block_bbox
    block_area = max(1.0, (bx1 - bx0) * (by1 - by0))

    for tx0, ty0, tx1, ty1 in table_bboxes:
        inter_x = max(0.0, min(bx1, tx1) - max(bx0, tx0))
        inter_y = max(0.0, min(by1, ty1) - max(by0, ty0))
        intersection = inter_x * inter_y
        if intersection / block_area >= overlap_threshold:
            return True

    return False

# ─────────────────────────────────────────────────────────────────────────────

def get_pdf_page_count(pdf_path: str) -> int:
    """Get PDF page count quickly with PyMuPDF, with pdfplumber fallback."""
    if fitz:
        pdf_doc = fitz.open(pdf_path)
        try:
            return len(pdf_doc)
        finally:
            pdf_doc.close()

    if pdfplumber:
        with pdfplumber.open(pdf_path) as pdf:
            return len(pdf.pages)

    return 0

def normalize_conversion_mode(mode: Optional[str]) -> str:
    normalized = (mode or DEFAULT_CONVERSION_MODE).strip().lower()
    if normalized in ('editable', 'no-ocr', 'non-ocr', 'no_ocr'):
        return 'no-ocr'
    if normalized == 'ocr':
        return 'ocr'
    if normalized == 'layout':
        return 'layout'
    return DEFAULT_CONVERSION_MODE

def check_ocr_runtime() -> Tuple[bool, str]:
    """Check whether OCR dependencies are available on the current machine."""
    if not pytesseract:
        return False, 'The Python package "pytesseract" is not installed on the server.'
    if not Image:
        return False, 'The Python package "Pillow" is not installed on the server.'
    if not fitz:
        return False, 'PyMuPDF is not installed on the server.'

    try:
        pytesseract.get_tesseract_version()
    except Exception as e:
        return False, f'Tesseract OCR is not available on the server: {e}'

    return True, ''

def convert_pdf_with_pdf2docx(pdf_path: str, output_path: str, page_count: int = 0) -> Tuple[bool, str]:
    """Use pdf2docx for a stronger editable DOCX conversion when available."""
    if Pdf2DocxConverter is None:
        return False, 'The pdf2docx Python package is not installed on the server.'

    import threading

    timeout_seconds = max(30, min(180, (page_count or 5) * 12))
    result_holder = {'success': False, 'message': '', 'error': None}

    def _run_conversion():
        converter = None
        try:
            settings = {
                'ignore_page_error': True,
                'multi_processing': False,
                'parse_lattice_table': True,
                'parse_stream_table': True,
                'extract_stream_table': page_count <= 200 if page_count else True,
                'delete_end_line_hyphen': False,
                'connected_border': True,
                'line_overlap_threshold': 0.9,
                'max_line_spacing_ratio': 2.0,
                'line_separate_threshold': 5,
                'min_section_height': 20,
            }

            converter = Pdf2DocxConverter(pdf_path)
            converter.convert(output_path, **settings)

            if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                result_holder['success'] = True
                result_holder['message'] = f'Enhanced editable DOCX created ({os.path.getsize(output_path)} bytes)'
            else:
                result_holder['message'] = 'pdf2docx completed but did not create a valid DOCX.'
        except Exception as e:
            result_holder['error'] = str(e)
            result_holder['message'] = str(e)
        finally:
            if converter is not None:
                try:
                    converter.close()
                except Exception:
                    pass

    logger.info("🧠 Step 1: Running pdf2docx conversion (timeout=%ds)...", timeout_seconds)

    thread = threading.Thread(target=_run_conversion, daemon=True)
    thread.start()
    thread.join(timeout=timeout_seconds)

    if thread.is_alive():
        logger.warning("pdf2docx conversion timed out after %d seconds, falling back to legacy.", timeout_seconds)
        try:
            if os.path.exists(output_path):
                os.remove(output_path)
        except Exception:
            pass
        return False, f'pdf2docx timed out after {timeout_seconds}s on this PDF.'

    if result_holder['success']:
        return True, result_holder['message']

    logger.warning(f"pdf2docx conversion failed: {result_holder['message']}")
    return False, result_holder['message']

def render_page_to_pil(page, scale: float = OCR_RENDER_SCALE):
    """Render a PyMuPDF page to a PIL image."""
    if not fitz or not Image:
        raise RuntimeError('OCR image rendering requires PyMuPDF and Pillow.')

    pix = page.get_pixmap(matrix=fitz.Matrix(scale, scale), alpha=False)
    image_bytes = pix.tobytes('png')
    return Image.open(io.BytesIO(image_bytes)).convert('RGB')

def extract_ocr_pages(pdf_path: str) -> List[dict]:
    """Run OCR on each PDF page and build editable text blocks."""
    available, reason = check_ocr_runtime()
    if not available:
        raise RuntimeError(reason)

    pages_content = []
    pdf_doc = fitz.open(pdf_path)

    try:
        for page_num, page in enumerate(pdf_doc):
            image = render_page_to_pil(page, OCR_RENDER_SCALE)
            ocr_data = pytesseract.image_to_data(
                image,
                output_type=pytesseract.Output.DICT,
                config='--oem 3 --psm 6',
            )

            page_rect = page.rect
            page_content = {
                'page': page_num,
                'width_in': page_rect.width / 72,
                'height_in': page_rect.height / 72,
                'blocks': [],
            }

            lines = {}
            total_items = len(ocr_data.get('text', []))
            for index in range(total_items):
                raw_text = ocr_data['text'][index]
                text = raw_text.strip()
                conf_raw = str(ocr_data.get('conf', [''])[index]).strip()

                if not text:
                    continue

                try:
                    confidence = float(conf_raw)
                except Exception:
                    confidence = 0.0

                if confidence < 25:
                    continue

                key = (
                    ocr_data.get('block_num', [0])[index],
                    ocr_data.get('par_num', [0])[index],
                    ocr_data.get('line_num', [0])[index],
                )

                if key not in lines:
                    lines[key] = {'words': [], 'tops': [], 'heights': []}

                lines[key]['words'].append((ocr_data.get('left', [0])[index], text))
                lines[key]['tops'].append(ocr_data.get('top', [0])[index])
                lines[key]['heights'].append(ocr_data.get('height', [0])[index])

            sorted_lines = sorted(
                lines.values(),
                key=lambda item: min(item['tops']) if item['tops'] else 0,
            )

            for line in sorted_lines:
                ordered_words = ' '.join(word for _, word in sorted(line['words'], key=lambda item: item[0])).strip()
                if not ordered_words:
                    continue

                avg_height = sum(line['heights']) / max(len(line['heights']), 1)
                estimated_font_size = max(10, min(22, int(round(avg_height / OCR_RENDER_SCALE))))
                is_heading = ordered_words.isupper() and len(ordered_words) > 4

                page_content['blocks'].append({
                    'type': 'text',
                    'data': {
                        'text': ordered_words,
                        'font_name': 'Calibri',
                        'font_size': max(estimated_font_size, 12 if is_heading else 10),
                        'is_bold': is_heading,
                        'is_italic': False,
                        'color': 0,
                    }
                })

            pages_content.append(page_content)
    finally:
        pdf_doc.close()

    return pages_content

def get_visual_render_scale(page_count: int) -> float:
    """Choose a render scale that balances fidelity with runtime for large files."""
    if page_count >= 1200:
        return 1.1
    if page_count >= 700:
        return 1.2
    if page_count >= 350:
        return 1.35
    if page_count >= 175:
        return 1.5
    return 1.7

def render_pdf_pages_for_docx(pdf_path: str, output_dir: str, page_count: int = 0) -> List[dict]:
    """Render each PDF page to an image so the DOCX preserves the original page layout."""
    rendered_pages = []

    if not fitz:
        logger.warning("PyMuPDF not available - cannot render PDF pages visually")
        return rendered_pages

    os.makedirs(output_dir, exist_ok=True)
    scale = get_visual_render_scale(page_count)

    pdf_doc = fitz.open(pdf_path)
    try:
        logger.info(f"Rendering {len(pdf_doc)} PDF pages for layout-preserving DOCX output at scale {scale:.2f}")

        for page_index, page in enumerate(pdf_doc):
            rect = page.rect
            pix = page.get_pixmap(matrix=fitz.Matrix(scale, scale), alpha=False)
            image_path = os.path.join(output_dir, f'page_{page_index + 1:04d}.png')
            pix.save(image_path)

            rendered_pages.append({
                'page': page_index,
                'image_path': image_path,
                'width_in': rect.width / 72,
                'height_in': rect.height / 72,
            })

        logger.info(f"Rendered {len(rendered_pages)} page images for DOCX output")
    finally:
        pdf_doc.close()

    return rendered_pages

def apply_section_page_layout(section, width_in: float, height_in: float, margin_in: float = 0.2) -> None:
    """Set the Word section size close to the original PDF page."""
    safe_width = max(width_in, 4.0)
    safe_height = max(height_in, 4.0)

    section.page_width = Inches(safe_width)
    section.page_height = Inches(safe_height)
    section.top_margin = Inches(margin_in)
    section.bottom_margin = Inches(margin_in)
    section.left_margin = Inches(margin_in)
    section.right_margin = Inches(margin_in)

def create_docx_from_page_images(output_path: str, rendered_pages: List[dict]) -> bool:
    """Create a DOCX where each Word page is a visual rendering of one PDF page."""
    if not rendered_pages:
        logger.error("No rendered page images were available for DOCX creation")
        return False

    try:
        doc = Document()

        first_page = rendered_pages[0]
        apply_section_page_layout(doc.sections[0], first_page['width_in'], first_page['height_in'])
        doc.core_properties.title = Path(output_path).stem

        for index, page_info in enumerate(rendered_pages):
            if index > 0:
                section = doc.add_section(WD_SECTION.NEW_PAGE)
                apply_section_page_layout(section, page_info['width_in'], page_info['height_in'])
            else:
                section = doc.sections[0]

            image_path = page_info['image_path']
            if not os.path.exists(image_path):
                logger.warning(f"Rendered page image missing: {image_path}")
                continue

            if doc.paragraphs:
                paragraph = doc.paragraphs[-1]
                if paragraph.text or paragraph.runs:
                    paragraph = doc.add_paragraph()
            else:
                paragraph = doc.add_paragraph()

            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1

            usable_width = section.page_width - section.left_margin - section.right_margin
            usable_height = section.page_height - section.top_margin - section.bottom_margin
            run = paragraph.add_run()
            run.add_picture(image_path, width=usable_width, height=usable_height)

        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)

        doc.save(output_path)
        return os.path.exists(output_path) and os.path.getsize(output_path) > 0
    except Exception as e:
        logger.error(f"Error creating layout-preserving DOCX: {e}")
        logger.exception("Layout-preserving DOCX creation failed")
        return False

def save_embedded_block_image(block: dict, output_dir: str, page_num: int, block_index: int) -> Optional[str]:
    """Save an image block extracted from PyMuPDF text dict output."""
    image_bytes = block.get('image')
    if not image_bytes:
        return None

    ext = (block.get('ext') or 'png').lower()
    if ext == 'jpg':
        ext = 'jpeg'

    image_path = os.path.join(output_dir, f'block_p{page_num}_b{block_index}.{ext}')
    try:
        os.makedirs(output_dir, exist_ok=True)
        with open(image_path, 'wb') as image_file:
            image_file.write(image_bytes)
        return image_path if os.path.exists(image_path) else None
    except Exception as e:
        logger.warning(f"Could not save embedded image block on page {page_num}: {e}")
        return None

# ─────────────────────────────────────────────────────────────────────────────
# TEXT EXTRACTION  (now skips blocks that overlap with table regions)
# ─────────────────────────────────────────────────────────────────────────────

def extract_text_and_layout(
    pdf_path: str,
    temp_dir: Optional[str] = None,
    table_bboxes_by_page: Optional[Dict[int, List[tuple]]] = None,
) -> List[dict]:
    """Extract text with layout information from PDF.

    Parameters
    ----------
    pdf_path:
        Path to the source PDF.
    temp_dir:
        Folder used to store formula/image snippets.
    table_bboxes_by_page:
        Mapping of page_num → list of table bounding boxes ``(x0,y0,x1,y1)``.
        Any PyMuPDF text block whose bbox substantially overlaps a table region
        is **skipped** here; the table will be rendered as a proper Word table
        by ``create_docx_from_extracted_content`` instead.  This prevents the
        classic bug where table cell text appears twice — once as flat
        paragraphs and once inside the table structure.
    """
    pages_content = []
    
    if not fitz:
        logger.warning("PyMuPDF not available - using pdfplumber for text")
        return extract_text_pdfplumber(pdf_path)
    
    try:
        pdf_doc = fitz.open(pdf_path)
        
        for page_num, page in enumerate(pdf_doc):
            text_dict = page.get_text("dict")
            blocks = text_dict.get("blocks", [])
            page_rect = page.rect
            
            page_content = {
                'page': page_num,
                'width_in': page_rect.width / 72,
                'height_in': page_rect.height / 72,
                'blocks': []
            }

            # Table bboxes for this specific page (may be empty list)
            page_table_bboxes = (table_bboxes_by_page or {}).get(page_num, [])
            
            for block_index, block in enumerate(blocks):
                if block['type'] == 0:  # Text block
                    # ── NEW: skip blocks whose content belongs to a table ────
                    block_bbox = block.get('bbox', (0, 0, 0, 0))
                    if page_table_bboxes and _block_is_inside_table(block_bbox, page_table_bboxes):
                        logger.debug(
                            f"Skipping text block (table region) on page {page_num}: "
                            f"bbox={block_bbox}"
                        )
                        continue
                    # ────────────────────────────────────────────────────────

                    for line in block.get("lines", []):
                        spans_data = []
                        for span in line.get("spans", []):
                            text_content = span.get('text', '')
                            if not text_content.strip():
                                continue
                            font_name_raw = span.get('font', 'Calibri')
                            color_int = span.get('color', 0)
                            spans_data.append({
                                'text': text_content,
                                'font_name': font_name_raw.split('+')[-1] if '+' in font_name_raw else font_name_raw,
                                'font_size': round(span.get('size', 11)),
                                'is_bold': 'Bold' in font_name_raw or 'bold' in font_name_raw.lower(),
                                'is_italic': 'Italic' in font_name_raw or 'italic' in font_name_raw.lower() or 'Oblique' in font_name_raw,
                                'color': color_int,
                            })

                        if not spans_data:
                            continue

                        complete_line = ''.join(s['text'] for s in spans_data).strip()
                        line_bbox = line.get('bbox') or block.get('bbox')

                        if not complete_line:
                            continue

                        if temp_dir and looks_like_formula(complete_line):
                            formula_image_path = save_formula_snippet(
                                pdf_doc, page, line_bbox, temp_dir,
                                page_num, len(page_content['blocks'])
                            )

                            if formula_image_path:
                                page_content['blocks'].append({
                                    'type': 'formula-image',
                                    'data': {
                                        'text': complete_line,
                                        'image_path': formula_image_path
                                    }
                                })
                                continue

                        page_content['blocks'].append({
                            'type': 'text',
                            'data': {
                                'text': complete_line,
                                'spans': spans_data,
                                'font_name': spans_data[0]['font_name'],
                                'font_size': spans_data[0]['font_size'],
                                'is_bold': spans_data[0]['is_bold'],
                                'is_italic': spans_data[0]['is_italic'],
                                'color': spans_data[0]['color'],
                            }
                        })

                elif block['type'] == 1 and temp_dir:
                    # ── Image block: only add if NOT inside a table region ──
                    block_bbox = block.get('bbox', (0, 0, 0, 0))
                    if page_table_bboxes and _block_is_inside_table(block_bbox, page_table_bboxes):
                        logger.debug(f"Skipping image block (table region) on page {page_num}")
                        continue

                    image_path = save_embedded_block_image(block, temp_dir, page_num, block_index)
                    bbox = block.get('bbox') or [0, 0, 0, 0]
                    width_pt = max(float(bbox[2]) - float(bbox[0]), 1.0)
                    height_pt = max(float(bbox[3]) - float(bbox[1]), 1.0)

                    if image_path:
                        page_content['blocks'].append({
                            'type': 'image',
                            'data': {
                                'image_path': image_path,
                                'width_in': width_pt / 72,
                                'height_in': height_pt / 72,
                            }
                        })
            
            pages_content.append(page_content)
        
        pdf_doc.close()
        logger.info(f"Extracted text from {len(pages_content)} pages")
        
    except Exception as e:
        logger.error(f"Error extracting text: {e}")
        import traceback
        logger.error(traceback.format_exc())
    
    return pages_content

# ─────────────────────────────────────────────────────────────────────────────

def extract_text_pdfplumber(pdf_path: str) -> List[dict]:
    """Fallback text extraction using pdfplumber"""
    pages_content = []
    
    if not pdfplumber:
        logger.error("No PDF text extraction library available")
        return pages_content
    
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                text = page.extract_text()
                pages_content.append({
                    'page': page_num,
                    'blocks': [{'type': 'text', 'data': {'text': text}}]
                })
    
    except Exception as e:
        logger.error(f"Error with pdfplumber: {e}")
    
    return pages_content

def _color_int_to_rgb(color_int: int) -> Optional[RGBColor]:
    """Convert a PyMuPDF integer color to a python-docx RGBColor."""
    if color_int is None or color_int <= 0:
        return None
    r = (color_int >> 16) & 0xFF
    g = (color_int >> 8) & 0xFF
    b = color_int & 0xFF
    if r == 0 and g == 0 and b == 0:
        return None
    return RGBColor(r, g, b)


def _estimate_text_height_pt(text: str, font_size: float, usable_width_pt: float) -> float:
    """Rough estimate of how many points of vertical space a paragraph consumes."""
    if not text:
        return 0.0
    avg_char_width = font_size * 0.5
    chars_per_line = max(1, int(usable_width_pt / avg_char_width))
    num_lines = max(1, -(-len(text) // chars_per_line))
    line_height = font_size * 1.08
    return num_lines * line_height


def _set_compact_cell_formatting(cell) -> None:
    """Remove excessive internal spacing from table cells."""
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        # FIX: Ensure consistent font size and prevent text wrapping issues
        for run in paragraph.runs:
            run.font.size = Pt(9)
            run.font.name = 'Calibri'
    
    # FIX: Set cell vertical alignment to top for better text positioning
    try:
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    except Exception:
        pass
    
    # FIX: Adjust cell margins for tighter layout
    try:
        tc = cell._tc
        tcPr = tc.tcPr
        if tcPr is None:
            tcPr = OxmlElement('w:tcPr')
            tc.insert(0, tcPr)
        
        # Set cell margins (left, right, top, bottom)
        tcMar = OxmlElement('w:tcMar')
        for margin_name, w_val in [('left', '50'), ('right', '50'), ('top', '10'), ('bottom', '10')]:
            margin = OxmlElement(f'w:{margin_name}')
            margin.set(qn('w:w'), w_val)
            margin.set(qn('w:type'), 'dxa')
            tcMar.append(margin)
        tcPr.append(tcMar)
    except Exception as e:
        logger.debug(f"Could not set cell margins: {e}")


def _set_table_borders(table) -> None:
    """Apply clean single-line borders to a table for professional appearance."""
    try:
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
        borders = OxmlElement('w:tblBorders')
        
        # FIX: Use stronger black borders for better visibility
        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            element = OxmlElement(f'w:{edge}')
            element.set(qn('w:val'), 'single')
            element.set(qn('w:sz'), '6')  # Slightly thicker borders
            element.set(qn('w:space'), '0')
            element.set(qn('w:color'), '000000')  # Black borders for clarity
            borders.append(element)
        
        tblPr.append(borders)
    except Exception as e:
        logger.warning(f"Could not set table borders: {e}")


def _apply_span_formatting(run, span_info: dict) -> None:
    """Apply font name, size, bold, italic, and color from a span dict to a docx Run."""
    font_name = span_info.get('font_name', 'Calibri')
    for prefix in ('AAAAAB+', 'AAAAAC+', 'AAAAAD+', 'AAAAAE+'):
        if font_name.startswith(prefix):
            font_name = font_name[len(prefix):]
    run.font.name = font_name
    run.font.size = Pt(span_info.get('font_size', 11))
    if span_info.get('is_bold'):
        run.font.bold = True
    if span_info.get('is_italic'):
        run.font.italic = True
    color_rgb = _color_int_to_rgb(span_info.get('color', 0))
    if color_rgb:
        run.font.color.rgb = color_rgb


def create_docx_from_extracted_content(
    output_path: str,
    pages_content: List[dict],
    tables: List[dict],
    images: dict
) -> bool:
    """Create a DOCX document from extracted PDF content."""

    try:
        logger.info("Creating new DOCX document...")
        doc = Document()

        try:
            style = doc.styles['Normal']
            style.font.name = 'Calibri'
            style.font.size = Pt(11)
            style.paragraph_format.space_before = Pt(0)
            style.paragraph_format.space_after = Pt(1)
            style.paragraph_format.line_spacing = 1.0
        except Exception as e:
            logger.warning(f"Could not set default style: {e}")

        if not pages_content:
            logger.warning("No page content to add")
            doc.add_paragraph("No text content could be extracted from the PDF.")

        content_count = 0
        MARGIN_IN = 0.35

        for page_idx, page_content in enumerate(pages_content):
            try:
                page_num = page_content.get('page', page_idx)
                logger.debug(f"Processing page {page_num}...")

                page_width_in = page_content.get('width_in', 8.27)
                page_height_in = page_content.get('height_in', 11.69)

                if page_idx == 0:
                    section = doc.sections[0]
                    apply_section_page_layout(section, page_width_in, page_height_in, margin_in=MARGIN_IN)
                else:
                    section = doc.add_section(WD_SECTION.NEW_PAGE)
                    apply_section_page_layout(section, page_width_in, page_height_in, margin_in=MARGIN_IN)

                usable_width_in = max((section.page_width - section.left_margin - section.right_margin) / 914400, 1.0)
                usable_width_pt = usable_width_in * 72
                usable_height_pt = max((section.page_height - section.top_margin - section.bottom_margin) / 914400, 1.0) * 72

                content_height_pt = 0.0
                OVERFLOW_SAFETY_MARGIN_PT = 14

                page_tables = [t for t in tables if t.get('page') == page_num]

                blocks = page_content.get('blocks', [])
                for block in blocks:
                    try:
                        if block.get('type') == 'text':
                            text_info = block.get('data', {})
                            text = text_info.get('text', '').strip()

                            if text:
                                font_size = text_info.get('font_size', 11)
                                est_height = _estimate_text_height_pt(text, font_size, usable_width_pt) + 2

                                if content_height_pt > 0 and (content_height_pt + est_height + OVERFLOW_SAFETY_MARGIN_PT) > usable_height_pt:
                                    section = doc.add_section(WD_SECTION.NEW_PAGE)
                                    apply_section_page_layout(section, page_width_in, page_height_in, margin_in=MARGIN_IN)
                                    content_height_pt = 0.0

                                p = doc.add_paragraph()

                                spans = text_info.get('spans', [])
                                if spans and len(spans) > 0:
                                    # FIX: Join all spans into complete text first, then apply primary formatting
                                    # This prevents text fragmentation
                                    complete_text = ''.join(s.get('text', '') for s in spans)
                                    complete_text = re.sub(r' +', ' ', complete_text).strip()
                                    
                                    if complete_text:
                                        run = p.add_run(complete_text)
                                        # Apply formatting from first span (primary formatting for the line)
                                        try:
                                            _apply_span_formatting(run, spans[0])
                                        except Exception as e:
                                            logger.warning(f"Could not apply span formatting: {e}")
                                else:
                                    run = p.add_run(text)
                                    try:
                                        _apply_span_formatting(run, text_info)
                                    except Exception as e:
                                        logger.warning(f"Could not apply text formatting: {e}")

                                p.paragraph_format.space_before = Pt(0)
                                p.paragraph_format.space_after = Pt(1)
                                p.paragraph_format.line_spacing = 1.0

                                content_height_pt += est_height
                                content_count += 1

                        elif block.get('type') == 'image':
                            image_info = block.get('data', {})
                            image_path = image_info.get('image_path')

                            if image_path and os.path.exists(image_path):
                                img_height_in = image_info.get('height_in') or 2.0
                                img_height_pt = img_height_in * 72

                                if content_height_pt > 0 and (content_height_pt + img_height_pt + OVERFLOW_SAFETY_MARGIN_PT) > usable_height_pt:
                                    section = doc.add_section(WD_SECTION.NEW_PAGE)
                                    apply_section_page_layout(section, page_width_in, page_height_in, margin_in=MARGIN_IN)
                                    content_height_pt = 0.0

                                paragraph = doc.add_paragraph()
                                paragraph.paragraph_format.space_before = Pt(1)
                                paragraph.paragraph_format.space_after = Pt(1)
                                run = paragraph.add_run()
                                requested_width = image_info.get('width_in') or usable_width_in
                                safe_width = min(max(requested_width, 0.6), usable_width_in)
                                run.add_picture(image_path, width=Inches(safe_width))
                                content_height_pt += img_height_pt + 2
                                content_count += 1

                        elif block.get('type') == 'formula-image':
                            formula_info = block.get('data', {})
                            image_path = formula_info.get('image_path')
                            formula_text = formula_info.get('text', '').strip()

                            if image_path and os.path.exists(image_path):
                                est_h = 36
                                if content_height_pt > 0 and (content_height_pt + est_h + OVERFLOW_SAFETY_MARGIN_PT) > usable_height_pt:
                                    section = doc.add_section(WD_SECTION.NEW_PAGE)
                                    apply_section_page_layout(section, page_width_in, page_height_in, margin_in=MARGIN_IN)
                                    content_height_pt = 0.0

                                paragraph = doc.add_paragraph()
                                paragraph.paragraph_format.space_before = Pt(1)
                                paragraph.paragraph_format.space_after = Pt(1)
                                run = paragraph.add_run()
                                run.add_picture(image_path, width=Inches(min(usable_width_in, 5.8)))
                                content_height_pt += est_h
                                content_count += 1
                            elif formula_text:
                                p = doc.add_paragraph(formula_text)
                                p.paragraph_format.space_before = Pt(0)
                                p.paragraph_format.space_after = Pt(1)
                                p.paragraph_format.line_spacing = 1.0
                                content_height_pt += 14
                                content_count += 1

                    except Exception as e:
                        logger.warning(f"Error processing text block: {e}")
                        continue

                # ── Render tables for this page as proper Word tables ────────
                for table_info in page_tables:
                    try:
                        table_data = table_info.get('data', [])
                        if not table_data or len(table_data) == 0:
                            continue

                        # Remove rows that are entirely None/empty
                        table_data = [
                            row for row in table_data
                            if any(cell is not None and str(cell).strip() for cell in row)
                        ]
                        if not table_data:
                            continue

                        est_table_h = len(table_data) * 18
                        if content_height_pt > 0 and (content_height_pt + est_table_h + OVERFLOW_SAFETY_MARGIN_PT) > usable_height_pt:
                            section = doc.add_section(WD_SECTION.NEW_PAGE)
                            apply_section_page_layout(section, page_width_in, page_height_in, margin_in=MARGIN_IN)
                            content_height_pt = 0.0

                        cols = max(len(row) for row in table_data) if table_data else 1
                        table = doc.add_table(rows=len(table_data), cols=cols)
                        try:
                            table.style = 'Table Grid'
                        except Exception:
                            pass
                        _set_table_borders(table)
                        table.alignment = WD_TABLE_ALIGNMENT.CENTER

                        for row_idx, row_data in enumerate(table_data):
                            for col_idx, cell_data in enumerate(row_data):
                                if col_idx < len(table.rows[row_idx].cells):
                                    cell = table.rows[row_idx].cells[col_idx]
                                    # FIX: Properly clean cell text to prevent fragmentation
                                    cell_text = str(cell_data) if cell_data is not None else ''
                                    # Remove excessive newlines but preserve intentional line breaks
                                    cell_text = re.sub(r'\n\s*\n', '\n', cell_text)
                                    # Remove leading/trailing whitespace from each line
                                    cell_lines = [line.strip() for line in cell_text.split('\n')]
                                    # Filter out empty lines
                                    cell_lines = [line for line in cell_lines if line]
                                    # FIX: For table cells, join multiple lines with space to prevent line breaks
                                    # This keeps cell content on a single line like iLovePDF
                                    if len(cell_lines) > 1:
                                        cell_text = ' '.join(cell_lines)
                                    elif len(cell_lines) == 1:
                                        cell_text = cell_lines[0]
                                    else:
                                        cell_text = ''
                                    
                                    cell.text = cell_text
                                    _set_compact_cell_formatting(cell)
                                    # Ensure text is visible
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.font.color.rgb = RGBColor(0, 0, 0)

                        content_height_pt += est_table_h
                        content_count += 1
                        logger.debug(f"Added table with {len(table_data)} rows on page {page_num}")
                    except Exception as e:
                        logger.warning(f"Could not add table on page {page_num}: {e}")
                        continue

            except Exception as e:
                logger.error(f"Error processing page {page_idx}: {e}")
                continue

        if content_count == 0:
            logger.warning("No content added to document")
            return False

        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)

        try:
            logger.info(f"Saving document to: {output_path}")
            doc.save(output_path)

            if os.path.exists(output_path):
                file_size = os.path.getsize(output_path)
                logger.info(f"Document saved successfully: {output_path} ({file_size} bytes)")
                return True
            else:
                logger.error(f"Document file was not created at {output_path}")
                return False

        except Exception as e:
            logger.error(f"Error saving document: {e}")
            import traceback
            logger.error(traceback.format_exc())
            return False

    except Exception as e:
        logger.error(f"Error creating DOCX: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return False

def convert_pdf_to_word(pdf_path: str, output_path: str, temp_dir: str = None, mode: str = DEFAULT_CONVERSION_MODE) -> Tuple[bool, str]:
    """Main conversion function with comprehensive error handling"""
    
    if not pdf_path or not os.path.exists(pdf_path):
        error_msg = f"PDF file not found: {pdf_path}"
        logger.error(error_msg)
        return False, error_msg
    
    if not output_path:
        error_msg = "Output path not specified"
        logger.error(error_msg)
        return False, error_msg
    
    if temp_dir is None:
        temp_dir = os.path.join(os.path.dirname(output_path), 'temp_images')
    
    try:
        os.makedirs(temp_dir, exist_ok=True)
    except Exception as e:
        error_msg = f"Could not create temp directory {temp_dir}: {e}"
        logger.error(error_msg)
        return False, error_msg
    
    conversion_mode = normalize_conversion_mode(mode)
    conversion_success = False
    
    try:
        logger.info(f"🚀 Starting PDF to Word conversion: {pdf_path}")
        logger.info(f"📁 Output path: {output_path}")
        logger.info(f"📦 Temp directory: {temp_dir}")

        page_count = get_pdf_page_count(pdf_path)
        logger.info(f"📄 Page count: {page_count}")
        logger.info(f"🧭 Conversion mode: {conversion_mode}")

        if conversion_mode == 'layout' and fitz:
            try:
                logger.info("🖼️ Step 1: Rendering each PDF page for layout-preserving DOCX output...")
                rendered_pages = render_pdf_pages_for_docx(pdf_path, temp_dir, page_count)

                if not rendered_pages:
                    error_msg = "The PDF pages could not be rendered for Word conversion."
                    logger.error(error_msg)
                    return False, error_msg

                logger.info("📄 Step 2: Building DOCX with one visual page per Word page...")
                success = create_docx_from_page_images(output_path, rendered_pages)

                if not success:
                    error_msg = "Failed to create the layout-preserving DOCX document."
                    logger.error(error_msg)
                    return False, error_msg

                conversion_success = True
            except Exception as e:
                error_msg = f"Layout-preserving PDF to Word conversion failed: {e}"
                logger.error(error_msg)
                logger.exception("Visual DOCX conversion failed")
                return False, error_msg

        elif conversion_mode == 'ocr':
            available, reason = check_ocr_runtime()
            if not available:
                return False, (
                    'OCR mode requires a server OCR setup. '
                    f'{reason} Install Tesseract OCR and the required Python packages to use OCR conversion.'
                )

            try:
                logger.info("🔎 Step 1: Running OCR on PDF pages...")
                pages_content = extract_ocr_pages(pdf_path)

                if not pages_content:
                    error_msg = 'OCR could not read any text from this PDF.'
                    logger.error(error_msg)
                    return False, error_msg

                logger.info("📄 Step 2: Building editable DOCX from OCR text...")
                success = create_docx_from_extracted_content(output_path, pages_content, [], {})

                if not success:
                    error_msg = 'Failed to create the OCR DOCX document.'
                    logger.error(error_msg)
                    return False, error_msg

                conversion_success = True
            except Exception as e:
                error_msg = f"OCR PDF to Word conversion failed: {e}"
                logger.error(error_msg)
                logger.exception("OCR DOCX conversion failed")
                return False, error_msg

        else:
            # ── no-ocr mode: try pdf2docx first, fall back to legacy ─────────
            if conversion_mode == 'layout' and not fitz:
                logger.warning("PyMuPDF is unavailable, falling back to no-OCR editable DOCX mode.")
            else:
                logger.info("Using no-OCR editable DOCX mode.")

            if Pdf2DocxConverter is not None:
                enhanced_success, enhanced_message = convert_pdf_with_pdf2docx(pdf_path, output_path, page_count)
                if enhanced_success:
                    logger.info(f"✓ {enhanced_message}")
                    conversion_success = True
                else:
                    logger.warning(
                        "Enhanced pdf2docx conversion did not finish cleanly, falling back "
                        f"to the legacy extractor. Reason: {enhanced_message}"
                    )
            else:
                logger.warning("pdf2docx is not available, using the legacy extractor.")

            if not conversion_success:
                fast_mode = page_count >= FAST_MODE_PAGE_THRESHOLD if page_count else False
                if fast_mode:
                    logger.warning(
                        f"Large PDF detected ({page_count} pages). Fast mode: table extraction skipped."
                    )

                # ── STEP A: Extract tables (with bboxes) ─────────────────
                if fast_mode:
                    tables = []
                    table_bboxes_by_page: Dict[int, List[tuple]] = {}
                    logger.warning("⚡ Fast mode: skipping table extraction for this very large PDF.")
                else:
                    try:
                        logger.info("📊 Step 2: Extracting tables with legacy fallback...")
                        tables = extract_tables_from_pdf(pdf_path)
                        # Build the page → bbox index used by text extraction
                        table_bboxes_by_page = _build_table_bbox_index(tables)
                        logger.info(f"✓ Extracted {len(tables)} tables across "
                                    f"{len(table_bboxes_by_page)} pages")
                    except Exception as e:
                        logger.warning(f"⚠️  Failed to extract tables: {e}")
                        tables = []
                        table_bboxes_by_page = {}

                # ── STEP B: Extract text, skipping table regions ──────────
                try:
                    logger.info("📝 Step 3: Extracting editable text, inline images, and formulas...")
                    pages_content = extract_text_and_layout(
                        pdf_path,
                        temp_dir,
                        table_bboxes_by_page,   # ← passes bbox filter
                    )

                    if not pages_content:
                        logger.warning("No text content extracted from PDF")
                        pages_content = []
                except Exception as e:
                    error_msg = f"Failed to extract text: {e}"
                    logger.error(error_msg)
                    return False, error_msg

                if not pages_content and not tables:
                    error_msg = "No content could be extracted from the PDF"
                    logger.error(error_msg)
                    return False, error_msg

                # ── STEP C: Build the DOCX ────────────────────────────────
                try:
                    logger.info("📄 Step 4: Creating editable DOCX document with legacy fallback...")
                    success = create_docx_from_extracted_content(
                        output_path,
                        pages_content,
                        tables,
                        {}
                    )

                    if not success:
                        error_msg = "Failed to create DOCX document - check logs for details"
                        logger.error(error_msg)
                        return False, error_msg

                    conversion_success = True

                except Exception as e:
                    error_msg = f"Error creating DOCX: {e}"
                    logger.error(error_msg)
                    return False, error_msg
        
        # Final verification
        if not os.path.exists(output_path):
            error_msg = "Output DOCX file was not created"
            logger.error(error_msg)
            return False, error_msg
        
        output_size = os.path.getsize(output_path)
        if output_size == 0:
            error_msg = "Output DOCX file is empty"
            logger.error(error_msg)
            return False, error_msg
        
        success_msg = f"✓ Successfully converted to {os.path.basename(output_path)} ({output_size} bytes)"
        logger.info(success_msg)
        return True, success_msg
    
    except Exception as e:
        error_msg = f"Unexpected error during conversion: {e}"
        logger.error(error_msg)
        import traceback
        logger.error(traceback.format_exc())
        return False, error_msg
    
    finally:
        if conversion_success:
            try:
                if os.path.exists(temp_dir):
                    shutil.rmtree(temp_dir, ignore_errors=True)
                logger.info("✓ Cleanup completed")
            except Exception as cleanup_err:
                logger.warning(f"Could not fully cleanup temp files: {cleanup_err}")

def main():
    """CLI entry point"""
    if len(sys.argv) < 3:
        print("Usage: python pdf_to_word_converter.py <input_pdf> <output_docx> [no-ocr|ocr|layout]")
        sys.exit(1)
    
    input_pdf = sys.argv[1]
    output_docx = sys.argv[2]
    conversion_mode = sys.argv[3] if len(sys.argv) > 3 else DEFAULT_CONVERSION_MODE
    
    success, message = convert_pdf_to_word(input_pdf, output_docx, mode=conversion_mode)
    
    result = {
        'success': success,
        'message': message,
        'input': input_pdf,
        'output': output_docx,
        'mode': normalize_conversion_mode(conversion_mode),
    }
    
    print(json.dumps(result))
    sys.exit(0 if success else 1)

if __name__ == '__main__':
    main()
