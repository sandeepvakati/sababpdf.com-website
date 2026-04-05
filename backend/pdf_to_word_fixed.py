#!/usr/bin/env python3
"""
Fixed PDF to Word Converter - Optimized for Table Conversion
Based on iLovePDF quality standards
Fixes: Text fragmentation, table cell splitting, line break issues
"""

import sys
import os
from pathlib import Path
import json
import logging
import re
import shutil
from typing import List, Tuple, Optional, Dict
from io import BytesIO

# Configure logging
LOG_LEVEL_NAME = os.environ.get('PDF_TO_WORD_LOG_LEVEL', 'WARNING').upper()
LOG_LEVEL = getattr(logging, LOG_LEVEL_NAME, logging.WARNING)
logging.basicConfig(
    level=LOG_LEVEL,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Try to import required packages
FITZ_AVAILABLE = False
try:
    import fitz  # PyMuPDF
    FITZ_AVAILABLE = True
except ImportError as e:
    logger.warning(f"PyMuPDF import error: {e}")
    pass

PDFPLUMBER_AVAILABLE = False
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError as e:
    logger.warning(f"pdfplumber import error: {e}")
    pass

DOCX_AVAILABLE = False
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm, Emu
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError as e:
    logger.warning(f"python-docx import error: {e}")
    pass

PIL_AVAILABLE = False
try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError as e:
    logger.warning(f"Pillow import error: {e}")
    pass

PDF2DOCX_AVAILABLE = False
try:
    from pdf2docx import Converter as Pdf2DocxConverter
    PDF2DOCX_AVAILABLE = True
except ImportError as e:
    logger.warning(f"pdf2docx import error: {e}")
    pass

# Constants
DEFAULT_DPI = 300
MARGIN_CM = 2.0


def check_dependencies():
    """Check and report available dependencies"""
    deps = {
        'PyMuPDF': FITZ_AVAILABLE,
        'pdfplumber': PDFPLUMBER_AVAILABLE,
        'python-docx': DOCX_AVAILABLE,
        'Pillow': PIL_AVAILABLE,
        'pdf2docx': PDF2DOCX_AVAILABLE,
    }
    missing = [name for name, available in deps.items() if not available]
    if missing:
        logger.warning(f"Missing optional dependencies: {', '.join(missing)}")
    return all(deps.values())


class FixedPDFToWordConverter:
    """
    Fixed PDF to Word converter with proper table and text handling.
    Key fixes:
    1. Proper text span joining - no fragmentation
    2. Table cells preserve complete text without line breaks
    3. Better table border formatting
    4. Intelligent content grouping
    """

    def __init__(self, pdf_path: str, output_path: str, mode: str = 'layout'):
        self.pdf_path = pdf_path
        self.output_path = output_path
        self.mode = mode
        self.temp_dir = os.path.join(os.path.dirname(output_path), 'temp_conversion')
        self.pdf_doc = None
        self.page_count = 0

    def open_pdf(self) -> bool:
        """Open PDF document with PyMuPDF"""
        if not FITZ_AVAILABLE:
            logger.error("PyMuPDF is required for high-quality conversion")
            return False
        try:
            self.pdf_doc = fitz.open(self.pdf_path)
            self.page_count = len(self.pdf_doc)
            logger.info(f"Opened PDF: {self.page_count} pages")
            return True
        except Exception as e:
            logger.error(f"Failed to open PDF: {e}")
            return False

    def close_pdf(self):
        """Close PDF document"""
        if self.pdf_doc:
            self.pdf_doc.close()
            self.pdf_doc = None

    def cleanup_temp(self):
        """Clean up temporary files"""
        if os.path.exists(self.temp_dir):
            try:
                shutil.rmtree(self.temp_dir, ignore_errors=True)
            except Exception as e:
                logger.warning(f"Could not clean temp dir: {e}")

    def get_page_dimensions(self, page_num: int) -> Tuple[float, float]:
        """Get page dimensions in inches"""
        page = self.pdf_doc[page_num]
        rect = page.rect
        return rect.width / 72, rect.height / 72

    def render_page_to_image(self, page_num: int, dpi: int = DEFAULT_DPI) -> Optional[Image.Image]:
        """Render a PDF page to a high-quality PIL Image"""
        if not FITZ_AVAILABLE or not PIL_AVAILABLE:
            return None
        try:
            page = self.pdf_doc[page_num]
            zoom = dpi / 72
            mat = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            img_data = pix.tobytes('png')
            img = Image.open(BytesIO(img_data))
            return img.convert('RGB')
        except Exception as e:
            logger.error(f"Failed to render page {page_num}: {e}")
            return None

    def extract_tables(self, page_num: int) -> List[Dict]:
        """
        Extract tables from PDF page using pdfplumber.
        KEY FIX: Properly handle table cell content without splitting.
        """
        if not PDFPLUMBER_AVAILABLE:
            return []

        tables = []
        try:
            with pdfplumber.open(self.pdf_path) as pdf:
                if page_num < len(pdf.pages):
                    page = pdf.pages[page_num]
                    page_tables = page.find_tables()

                    for table_idx, table in enumerate(page_tables):
                        table_data = table.extract()
                        if table_data and len(table_data) > 0:
                            # KEY FIX: Clean and join cell content properly
                            cleaned_data = []
                            for row in table_data:
                                cleaned_row = []
                                for cell in row:
                                    # FIX: Join fragmented text, remove excessive newlines
                                    cell_text = str(cell) if cell is not None else ''
                                    # Remove excessive line breaks but preserve intentional ones
                                    cell_text = re.sub(r'\n\s*\n', '\n', cell_text)
                                    # Remove leading/trailing whitespace from each line
                                    cell_lines = [line.strip() for line in cell_text.split('\n')]
                                    # Filter out empty lines
                                    cell_lines = [line for line in cell_lines if line]
                                    # Join with single space if multiple lines, or keep as single line
                                    if len(cell_lines) > 1:
                                        # For table cells, join with space to prevent line breaks
                                        cell_text = ' '.join(cell_lines)
                                    elif len(cell_lines) == 1:
                                        cell_text = cell_lines[0]
                                    else:
                                        cell_text = ''
                                    cleaned_row.append(cell_text)
                                cleaned_data.append(cleaned_row)

                            tables.append({
                                'type': 'table',
                                'data': cleaned_data,
                                'bbox': table.bbox,
                                'index': table_idx,
                                'num_rows': len(cleaned_data),
                                'num_cols': max(len(row) for row in cleaned_data) if cleaned_data else 0,
                            })

                    logger.info(f"Extracted {len(tables)} tables from page {page_num}")

        except Exception as e:
            logger.warning(f"Error extracting tables from page {page_num}: {e}")

        return tables

    def extract_text_blocks_fixed(self, page_num: int) -> List[Dict]:
        """
        Extract text blocks with FIXED span joining.
        KEY FIX: Join all spans in a line BEFORE adding to document.
        """
        if not FITZ_AVAILABLE:
            return []

        try:
            page = self.pdf_doc[page_num]
            text_dict = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)

            blocks = []
            for block in text_dict.get("blocks", []):
                if block.get("type") == 0:  # Text block
                    for line in block.get("lines", []):
                        # KEY FIX: Collect ALL spans first, then join them
                        line_text_parts = []
                        line_formatting = None
                        
                        for span in line.get("spans", []):
                            text = span.get("text", "")
                            if text.strip():
                                line_text_parts.append(text)
                                
                                # Capture formatting from first significant span
                                if line_formatting is None and len(text.strip()) > 0:
                                    line_formatting = {
                                        'font': span.get("font", "Calibri"),
                                        'size': span.get("size", 11),
                                        'color': span.get("color", 0),
                                        'bold': 'bold' in span.get("font", "").lower(),
                                        'italic': 'italic' in span.get("font", "").lower() or
                                                 'oblique' in span.get("font", "").lower(),
                                    }
                        
                        # KEY FIX: Join all parts into complete line text
                        if line_text_parts:
                            # Join with proper spacing (remove duplicate spaces)
                            complete_line = ''.join(line_text_parts)
                            # Clean up excessive whitespace but preserve structure
                            complete_line = re.sub(r' +', ' ', complete_line).strip()
                            
                            if complete_line:
                                blocks.append({
                                    'type': 'text',
                                    'text': complete_line,
                                    'formatting': line_formatting,
                                    'bbox': line.get("bbox"),
                                })

            return blocks
        except Exception as e:
            logger.error(f"Error extracting text from page {page_num}: {e}")
            return []

    def create_layout_document(self) -> bool:
        """
        Create a Word document with layout preservation (iLovePDF style).
        Each PDF page is rendered as an image - BEST for complex tables.
        """
        if not DOCX_AVAILABLE:
            logger.error("python-docx not available")
            return False

        try:
            doc = Document()
            style = doc.styles['Normal']
            style.font.name = 'Calibri'
            style.font.size = Pt(11)

            logger.info("Creating layout-preserving Word document...")

            for page_num in range(self.page_count):
                logger.info(f"Processing page {page_num + 1}/{self.page_count}")

                page_width_in, page_height_in = self.get_page_dimensions(page_num)

                if page_num > 0:
                    section = doc.add_section(WD_SECTION.NEW_PAGE)
                else:
                    section = doc.sections[0]

                section.page_width = Inches(page_width_in)
                section.page_height = Inches(page_height_in)
                section.top_margin = Cm(MARGIN_CM)
                section.bottom_margin = Cm(MARGIN_CM)
                section.left_margin = Cm(MARGIN_CM)
                section.right_margin = Cm(MARGIN_CM)

                img = self.render_page_to_image(page_num, dpi=DEFAULT_DPI)

                if img:
                    img_path = os.path.join(self.temp_dir, f'page_{page_num:04d}.png')
                    img.save(img_path, 'PNG', dpi=(DEFAULT_DPI, DEFAULT_DPI))

                    usable_width = section.page_width - section.left_margin - section.right_margin
                    usable_height = section.page_height - section.top_margin - section.bottom_margin

                    paragraph = doc.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run()

                    img_width_in = img.width / DEFAULT_DPI
                    img_height_in = img.height / DEFAULT_DPI
                    scale = min(usable_width.inches / img_width_in, usable_height.inches / img_height_in)

                    final_width = Inches(img_width_in * scale)
                    final_height = Inches(img_height_in * scale)

                    run.add_picture(img_path, width=final_width, height=final_height)
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.0

            doc.save(self.output_path)

            if os.path.exists(self.output_path):
                file_size = os.path.getsize(self.output_path)
                logger.info(f"Layout document saved: {file_size} bytes")
                return True

            return False

        except Exception as e:
            logger.error(f"Error creating layout document: {e}")
            import traceback
            logger.error(traceback.format_exc())
            return False

    def create_editable_document_fixed(self) -> bool:
        """
        Create an editable Word document with FIXED text and table handling.
        KEY FIXES:
        1. Join text spans properly (no fragmentation)
        2. Table cells contain complete text (no line breaks within cells)
        3. Proper table formatting with borders
        """
        if not DOCX_AVAILABLE:
            logger.error("python-docx not available")
            return False

        try:
            doc = Document()
            style = doc.styles['Normal']
            style.font.name = 'Calibri'
            style.font.size = Pt(11)
            style.paragraph_format.space_before = Pt(0)
            style.paragraph_format.space_after = Pt(3)
            style.paragraph_format.line_spacing = 1.15

            logger.info("Creating editable Word document with fixed table handling...")

            for page_num in range(self.page_count):
                logger.info(f"Processing page {page_num + 1}/{self.page_count}")

                page_width_in, page_height_in = self.get_page_dimensions(page_num)

                if page_num > 0:
                    section = doc.add_section(WD_SECTION.NEW_PAGE)
                else:
                    section = doc.sections[0]

                section.page_width = Inches(page_width_in)
                section.page_height = Inches(page_height_in)
                section.top_margin = Cm(MARGIN_CM)
                section.bottom_margin = Cm(MARGIN_CM)
                section.left_margin = Cm(MARGIN_CM)
                section.right_margin = Cm(MARGIN_CM)

                # Extract content with FIXED methods
                text_blocks = self.extract_text_blocks_fixed(page_num)
                tables = self.extract_tables(page_num)

                # Add text blocks with PROPER span joining
                for block in text_blocks:
                    if block['type'] == 'text':
                        paragraph = doc.add_paragraph()
                        
                        # KEY FIX: Add complete line text, not fragmented spans
                        run = paragraph.add_run(block['text'])
                        
                        # Apply formatting
                        formatting = block.get('formatting', {})
                        font = run.font
                        font.name = formatting.get('font', 'Calibri')
                        font.size = Pt(formatting.get('size', 11))
                        font.bold = formatting.get('bold', False)
                        font.italic = formatting.get('italic', False)

                        color = formatting.get('color', 0)
                        if color and color != 0:
                            r = (color >> 16) & 0xFF
                            g = (color >> 8) & 0xFF
                            b = color & 0xFF
                            if r != 0 or g != 0 or b != 0:
                                font.color.rgb = RGBColor(r, g, b)

                # Add tables with IMPROVED cell handling
                for table_info in tables:
                    table_data = table_info['data']
                    if table_data:
                        rows = len(table_data)
                        cols = table_info.get('num_cols', max(len(row) for row in table_data))

                        table = doc.add_table(rows=rows, cols=cols)
                        table.style = 'Table Grid'
                        table.alignment = WD_TABLE_ALIGNMENT.CENTER
                        
                        # Apply table borders
                        self._set_table_borders(table)

                        for row_idx, row_data in enumerate(table_data):
                            row = table.rows[row_idx]
                            for col_idx, cell_text in enumerate(row_data):
                                if col_idx < len(row.cells):
                                    cell = row.cells[col_idx]
                                    cell.text = cell_text
                                    
                                    # KEY FIX: Compact cell formatting
                                    # Prevent line breaks within cells
                                    for paragraph in cell.paragraphs:
                                        paragraph.paragraph_format.space_before = Pt(0)
                                        paragraph.paragraph_format.space_after = Pt(0)
                                        paragraph.paragraph_format.line_spacing = 1.0
                                        for run in paragraph.runs:
                                            run.font.size = Pt(9)
                                            run.font.color.rgb = RGBColor(0, 0, 0)

            doc.save(self.output_path)

            if os.path.exists(self.output_path):
                file_size = os.path.getsize(self.output_path)
                logger.info(f"Editable document saved: {file_size} bytes")
                return True

            return False

        except Exception as e:
            logger.error(f"Error creating editable document: {e}")
            import traceback
            logger.error(traceback.format_exc())
            return False

    def _set_table_borders(self, table):
        """Apply clean single-line borders to table"""
        try:
            tbl = table._tbl
            tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
            borders = OxmlElement('w:tblBorders')
            
            for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                element = OxmlElement(f'w:{edge}')
                element.set(qn('w:val'), 'single')
                element.set(qn('w:sz'), '4')
                element.set(qn('w:space'), '0')
                element.set(qn('w:color'), '000000')
                borders.append(element)
            
            tblPr.append(borders)
        except Exception as e:
            logger.warning(f"Could not set table borders: {e}")

    def convert(self) -> Tuple[bool, str]:
        """Main conversion method"""
        if not os.path.exists(self.pdf_path):
            return False, f"PDF file not found: {self.pdf_path}"

        if not check_dependencies():
            return False, "Missing required dependencies"

        if not self.open_pdf():
            return False, "Failed to open PDF file"

        try:
            os.makedirs(self.temp_dir, exist_ok=True)

            # Try pdf2docx first if available (best quality)
            if self.mode == 'no-ocr' and PDF2DOCX_AVAILABLE:
                try:
                    logger.info("Trying pdf2docx for best quality...")
                    converter = Pdf2DocxConverter(self.pdf_path)
                    converter.convert(self.output_path)
                    converter.close()

                    if os.path.exists(self.output_path) and os.path.getsize(self.output_path) > 0:
                        return True, "Successfully converted with pdf2docx (highest quality)"
                except Exception as e:
                    logger.warning(f"pdf2docx failed, using fallback: {e}")

            # Use layout mode for best visual fidelity (like iLovePDF)
            if self.mode == 'layout':
                logger.info("Using LAYOUT mode - best for tables and complex content")
                success = self.create_layout_document()
                if success:
                    return True, "Successfully converted (layout preserved)"
                else:
                    return False, "Layout conversion failed"

            # Use fixed editable mode
            else:
                logger.info("Using EDITABLE mode with fixed table handling")
                success = self.create_editable_document_fixed()
                if success:
                    return True, "Successfully converted (editable text)"
                else:
                    return False, "Editable conversion failed"

        except Exception as e:
            logger.error(f"Conversion error: {e}")
            import traceback
            logger.error(traceback.format_exc())
            return False, f"Conversion error: {e}"

        finally:
            self.close_pdf()
            self.cleanup_temp()


def convert_pdf_to_word(pdf_path: str, output_path: str, mode: str = 'layout') -> Tuple[bool, str]:
    """High-level conversion function"""
    converter = FixedPDFToWordConverter(pdf_path, output_path, mode)
    return converter.convert()


def main():
    """CLI entry point"""
    if len(sys.argv) < 3:
        print("Usage: python pdf_to_word_fixed.py <input_pdf> <output_docx> [layout|no-ocr]")
        sys.exit(1)

    input_pdf = sys.argv[1]
    output_docx = sys.argv[2]
    mode = sys.argv[3] if len(sys.argv) > 3 else 'layout'

    success, message = convert_pdf_to_word(input_pdf, output_docx, mode)

    result = {
        'success': success,
        'message': message,
        'input': input_pdf,
        'output': output_docx,
        'mode': mode,
    }

    print(json.dumps(result, indent=2))
    sys.exit(0 if success else 1)


if __name__ == '__main__':
    main()
